"""Модуль содержит процедуру отвечающую за создания отчета
в формате .docx по заданному экземпляру объекта класса
Hydraulic.Morfostvor.
"""

import os
import sys
from pathlib import Path

import numpy as np
from docx import Document
from pathvalidate import sanitize_filename

import hydraulic.config as config
from hydraulic.doc_lib import (insert_df_to_table, insert_figure,
                               insert_page_break)
from hydraulic.lib import rmdir, text_sanitize


def generate_morfostvor_report(morfostvor, out_filename, r=False):
    """Процедура создает отчет по заданному морфоствору и сохраняет
    его в указанный .docx файл.

    Args:
        morfostvor (hydraulic.Morfostvor): Экземпляр объекта класса Morfostvor
        out_filename (_type_): путь и названия файла куда будет сохранен отчет,
        на конце должно быть указание расширения .docx
        r (bool, optional): Параметр позволяет включить перезапись файла отчета,
        если выключен то отчет будет добавлен в конец документа. Defaults to False.
    """
    print("\n\nФормируем doc файл: ")
    doc_file = out_filename
    template_file = Path("hydraulic/assets/report_template.docx")

    # Создаем временную папку, и папку для графики если они не существуют
    temp_dir = Path(config.TEMP_DIR_NAME)
    temp_dir.mkdir(parents=True, exist_ok=True)

    # Создаем папку для сохранения отдельных изображений
    if config.PROFILE_SAVE_PICTURES:
        picture_dir = Path(
            str(Path(out_filename).parents[0]) + "/" + config.GRAPHICS_DIR_NAME
        )
        picture_dir.mkdir(parents=True, exist_ok=True)

    if r:
        doc = Document(template_file)
    else:
        if os.path.isfile(doc_file):
            doc = Document(doc_file)
            insert_page_break(doc)
        else:
            if config.REWRITE_DOC_FILE:
                print(
                    "    — Включена перезапись файла, удаляем старый и создаём новый."
                )
            else:
                print("    — Файл не найден! Создаём новый.")
            doc = Document(template_file)

    # Отрисовка смоченного периметра
    if config.PROFILE_WET_PERIMETER:
        morfostvor.fig_profile.draw_wet_perimeter()

    # Отрисовка верхней границы сооружения
    if morfostvor.top_limit:
        morfostvor.fig_profile.draw_top_limit(
            morfostvor.top_limit, text=morfostvor.top_limit_description
        )

    # Отрисовка границы предельного размыва профиля
    if morfostvor.erosion_limit and len(morfostvor.erosion_limit_coord) == 2:
        morfostvor.fig_profile.draw_erosion_limit(
            morfostvor.erosion_limit,
            morfostvor.erosion_limit_coord[0],
            morfostvor.erosion_limit_coord[1])
    elif morfostvor.erosion_limit and len(morfostvor.erosion_limit_coord) == 4:
        morfostvor.fig_profile.draw_erosion_limit(
            morfostvor.erosion_limit,
            morfostvor.erosion_limit_coord[0],
            morfostvor.erosion_limit_coord[1],
            morfostvor.erosion_limit_coord[2],
            morfostvor.erosion_limit_coord[3])
    elif morfostvor.erosion_limit:
        morfostvor.fig_profile.draw_erosion_limit(morfostvor.erosion_limit)

    # Отрисовка расчетных уровней воды на графике профиля
    morfostvor.fig_profile.draw_levels_on_profile(morfostvor.levels_result)
    morfostvor.fig_profile._update_limit()

    # TODO: сделать отрисовку линий урезов воды по каждому
    # участку УВ из описания ситуации исходного файла
    # Отрисовка урез воды на графике профиля
    if morfostvor.waterline and type(morfostvor.waterline) != str:
        morfostvor.fig_profile.draw_waterline(
            round(morfostvor.waterline, 2), color="blue", linestyle="-"
        )

    # Вставляем заголовок профиля
    doc.add_paragraph(morfostvor.title, style="З-приложение-подзаголовок")
    # Добавляем изображения профиля и гидравлической кривой
    print("    — Вставляем графику (профиль)... ", end="")
    insert_figure(doc, morfostvor.fig_profile.fig, width=16)

    # Подпись рисунков
    if config.GRAPHICS_TITLES_TEXT:
        doc.add_paragraph(
            f"{config.STRING['figure']}{morfostvor.fig_profile.morfostvor.title}",
            style="Р-название",
        )
    print("успешно!")

    if config.HYDRAULIC_CURVE:
        print("    — Вставляем графику (кривая QH)... ", end="")
        insert_figure(
            doc,
            morfostvor.fig_QH.fig,
            width=16
        )

        if config.GRAPHICS_TITLES_TEXT:
            doc.add_paragraph(
                f"{config.STRING['figure']}{morfostvor.fig_QH._ax_title_text}",
                style="Р-название",
            )
        print("успешно!")

    if config.HYDRAULIC_AND_SPEED_CURVE:
        print("    — Вставляем графику (кривая QHV)... ", end="")
        insert_figure(doc, morfostvor.fig_QHV.fig, width=16)

        if config.GRAPHICS_TITLES_TEXT:
            doc.add_paragraph(
                f"{config.STRING['figure']}{morfostvor.fig_QHV._ax_title_text}",
                style="Р-название",
            )
        print("успешно!")

    if config.SPEED_CURVE:
        print("    — Вставляем график кривой скоростей QV ... ", end="")
        insert_figure(doc, morfostvor.fig_QV.fig)
        print("успешно!")

        if config.GRAPHICS_TITLES_TEXT:
            doc.add_paragraph(
                f"{config.STRING['figure']}{morfostvor.fig_QV._ax_title_text}",
                style="Р-название",
            )

    if config.SPEED_VH_CURVE:
        print("    — Вставляем график кривой скоростей VH ... ", end="")
        insert_figure(doc, morfostvor.fig_VH.fig)
        print("успешно!")

        if config.GRAPHICS_TITLES_TEXT:
            doc.add_paragraph(
                f"{config.STRING['figure']}{morfostvor.fig_VH._ax_title_text}",
                style="Р-название",
            )

    if config.AREA_CURVE:
        print("    — Вставляем график кривой площадей ... ", end="")
        insert_figure(doc, morfostvor.fig_QF.fig)
        print("успешно!")

        if config.GRAPHICS_TITLES_TEXT:
            doc.add_paragraph(
                f"{config.STRING['figure']}{morfostvor.fig_QF._ax_title_text}",
                style="Р-название",
            )

    # Проверяем имя файла
    profile_name = sanitize_filename(morfostvor.title)

    # Сохраняем картинки в отдельные файлы в папку graphics
    if config.PROFILE_SAVE_PICTURES:
        morfostvor.fig_profile.fig.savefig(
            Path(f"{picture_dir}/{profile_name}.png", dpi=config.FIG_DPI)
        )
    if config.CURVE_SAVE_PICTURES:
        if config.HYDRAULIC_CURVE:
            morfostvor.fig_QH.fig.savefig(
                Path(f"{picture_dir}/{profile_name}_QH.png", dpi=config.FIG_DPI)
            )
        if config.HYDRAULIC_AND_SPEED_CURVE:
            morfostvor.fig_QHV.fig.savefig(
                Path(f"{picture_dir}/{profile_name}_QHV.png", dpi=config.FIG_DPI)
            )
        if config.SPEED_CURVE:
            morfostvor.fig_QV.fig.savefig(
                Path(f"{picture_dir}/{profile_name}_QV.png", dpi=config.FIG_DPI)
            )
        if config.AREA_CURVE:
            morfostvor.fig_QF.fig.savefig(
                Path(f"{picture_dir}/{profile_name}_QF.png", dpi=config.FIG_DPI)
            )

    # Вывод таблицы расчётных уровней воды
    print("    — Записываем таблицу уровней воды ... ", end="")
    insert_df_to_table(
        doc,
        morfostvor.levels_result[["P", "Q", "H"]],
        f"{config.STRING['table']}Расчётные уровни {morfostvor.strings['type']}",
        col_names=(
            "Обеспеченность P, %",
            "Расход Q, м³/сек",
            f"Уровень H, м{config.ALTITUDE_SYSTEM}",
        ),
        col_widths=(6, 6, 6),
        col_format=(":g", ":g", ":.2f"),
    )
    print("успешно!")

    # Вывод таблицы участков
    print("    — Записываем таблицу участков ... ", end="")
    # Заменяем пустые значения на прочерк и добавляем номер участка
    df_sectors = morfostvor.sectors_result.replace(np.NaN, '-')
    df_sectors.insert(loc=0, column='N', value=df_sectors.index + 1)

    prob_text = text_sanitize(
        morfostvor.probability[morfostvor.design_water_level_index][0],
        num_suffix="% обеспеченности",
    )

    insert_df_to_table(
        doc,
        df_sectors,
        f"{config.STRING['table']}Расчётные участки и их параметры",
        col_names=(
            "№",
            "Описание",
            "Уклон i, ‰",
            "Коэффициент шероховатости n",
            "Q при РУВВ, м³/сек",
            f"Hср при РУВВ, м{config.ALTITUDE_SYSTEM}",
            "Vср при РУВВ, м/сек",
            "B при РУВВ, м",
            "F при РУВВ, м²",
        ),
        col_widths=(1.3, 4, 4, 4, 4, 4, 4, 4, 4),
        col_format=(":d", "", ":g", ":.3f", ":.2f", ":.2f", ":.2f", ":.2f", ":.2f"),
        footer_text=(
            "Примечание: Расчетный уровень высоких вод (РУВВ) "
            f"принят по расходу {prob_text}."
        ),
    )
    print("успешно!")

    # Вывод таблицы гидравлической кривой
    print("    — Записываем таблицу кривой расхода воды ... ", end="")

    table = morfostvor.hydraulic_table.reset_index(0).loc["Сумма"].reset_index(drop=True)
    table_round = table.round(3)  # Округляем
    # Убираем столбец с коэффициентами Шези
    table_round = table_round.drop(columns=["Shezi"])

    if config.DOC_TABLE_SHORT:
        # Количество строк в таблице
        table_quant = table_round["УВ"].count()

        # Уменьшаем количество выводимых строк в таблицу
        # чтобы поместилось на один лист
        if table_quant <= 25:
            divider = 1
        elif table_quant > 25 and table_quant <= 50:
            divider = 2
        elif table_quant > 50 and table_quant <= 75:
            divider = 3
        elif table_quant > 75 and table_quant <= 100:
            divider = 4
        elif table_quant > 100 and table_quant <= 125:
            divider = 5
        else:
            divider = 10
    else:
        divider = 1

    # Записываем только чётные элементы таблицы
    table_round = table_round[table_round.index % divider == 0]

    insert_df_to_table(
        doc,
        table_round,
        f"{config.STRING['table']}Параметры расчёта кривой расхода {morfostvor.strings['type']}",
        col_names=(
            f"Отм. уровня H, м{config.ALTITUDE_SYSTEM}",
            "Площадь F, м²",
            "Ширина B, м",
            "Средняя глубина Hср, м",
            "Макс. глубина Hмакс, м",
            "Средняя скорость Vср, м/сек",
            "Расход Q, м³/сек",
        ),
        col_widths=(5, 5, 5, 5, 5, 5, 5),
        col_format=(":.2f", ":.3f", ":.3f", ":.3f", ":.3f", ":.3f", ":.3f"),
        footer_text=(
            f"Расчётный шаг: {morfostvor.dH:g} см. "
            f"В таблице приведён каждый {divider}-й результат расчёта."
        ),
    )

    print("успешно!")

    # Удаляем объект профиля
    morfostvor.fig_profile.clean()

    try:
        doc.save(doc_file)
    except PermissionError:
        print(
            "\nОшибка! Не удалось сохранить файл. "
            "Проверьте возможность записи файла по указанному пути."
        )
        print("Возможно записываемый файл уже существует и открыт.")
        sys.exit(1)

    # Удаляем временную папку со всем содержимым
    print("    — Удаляем временную папку ... ", end="")
    rmdir(Path(f"{config.TEMP_DIR_NAME}"))
    print("успешно!")
