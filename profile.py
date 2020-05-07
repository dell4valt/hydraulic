# -*- coding: utf-8 -*-
import sys
import os
import re
from dataclasses import dataclass, field

# matplotlib.use("TkAgg")
import matplotlib.pyplot as plt
from matplotlib import gridspec
from matplotlib.patches import Polygon
import matplotlib.patheffects as path_effects
import numpy as np
import pandas as pd
import scipy.interpolate as interpolate

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Cm

import xlrd

import config
from lib import *


# Скрываем предуприждение matplotlib
import warnings
import matplotlib.cbook
# warnings.filterwarnings("ignore",category=matplotlib.cbook.mplDeprecation)


@dataclass
class ProfileSector(object):
    """ Класс участка профиля (пойма, русло и т.д.)

        :param id: Номер участка
        :param name: Описание (название) участка
        :param start_point: Номер первой точки участка
        :param end_point: Номер последней точки участка
        :param roughness: Коэффициент шероховатости n
        :param slope: Уклон данного учаcтка I, ‰
        :param coord: Список с двумя подсписками координат x и y участка    
    """
    id: int
    name: str
    start_point: int
    end_point: int
    roughness: float
    slope: float
    coord: tuple

    def __post_init__(self):
        self.color = self.get_color()

    def get_color(self):
        name = self.name
        channel = re.findall('русло', name, flags=re.IGNORECASE)
        protoka = re.findall('протока', name, flags=re.IGNORECASE)
        floodplain = re.findall('пойма', name, flags=re.IGNORECASE)

        if channel:
            color = [0, .5, 1]
        elif protoka:
            color = [0, np.random.uniform(0, .5), np.random.uniform(.5, 1)]
        elif floodplain:
            color = [np.random.uniform(.3, 1), 0, 0]
        else:
            color = [np.random.uniform(0, 1), np.random.uniform(
                0, 1), np.random.uniform(0, 1)]
        return color

    def get_length(self):
        return round(self.coord[0][-1] - self.coord[0][0], 3)


@dataclass
class WaterSection(object):
    """ Класс водного сечения

        :param x: Точки x всего профиля
        :param y: Точки y всего профиля
        :param water_level: Уровнь воды
        :param water_section_x: Точки x водного сечения
        :param water_section_y: Точки y водного сечения
        :param width: Ширина водного сечения
        :param area: Площадь водного сечения
        :param average_depth: Средняя глубина
        :param max_depth: Максимальная глубина
        :param wet_perimeter: Смоченный периметер
        :pararm r_gidraulic: Гидравлический радиус
        :param start_point: Точка начала расчёта [point_index, y]

    """
    x: float
    y: float
    water_level: float
    water_section_x: list = field(default_factory=list)
    water_section_y: list = field(default_factory=list)
    width: float = 0.0
    area: float = 0.0
    average_depth: float = 0.0
    max_depth: float = 0.0
    wet_perimeter: float = 0.0
    r_gidraulic: float = 0.0
    start_point: list = field(default_factory=list)

    def __post_init__(self):
        # start_point=[self.y.index(min(self.y)), min(self.y)]
        boundry = self.boundry()
        if len(boundry) > 1:
            for water_boundry in boundry:
                try:
                    self._calculate_parameters(water_boundry)
                except IndexError:
                    print(
                        'Ошибка в определении границ урезов! Программа будет завершена.')
                    sys.exit(2)

            # Вычисления если урезов несколько
            self.width = sum(self.width)
            self.area = sum(self.area)
            self.average_depth = np.average(self.average_depth)
            self.max_depth = max(self.max_depth)
            self.wet_perimeter = sum(self.wet_perimeter)
            self.r_gidraulic = sum(self.r_gidraulic)

        else:
            try:
                self._calculate_parameters(boundry[0])
            except IndexError:
                print('Ошибка в определении границ урезов! Программа будет завершена.')
                sys.exit(2)

    def boundry(self):
        x = self.x
        y = self.y
        water_level = self.water_level  # Отметка уреза воды
        water_boundry_x, water_boundry_y, water_boundry_points = [], [], []
        result = []
        start_point = self.start_point

        if not start_point:
            start_point = [y.index(min(y)), min(y)]

        # Проверка на ошибку расположения уреза под поверхностью дна
        if water_level < min(y):
            print(
                'Ошибка! Уровень воды ниже низжей точки дна. Программа будет завершена с ошибкой.')
            sys.exit(1)
        else:
            # Цикл влево от стартовой точки
            for i in range(start_point[0], -1, -1):
                # Если индекс минимальной отметки совпадает с левой правой участка
                if start_point[0] == 0 and y[start_point[0]] <= water_level:
                    water_boundry_x.append(x[0])
                    water_boundry_y.append(water_level)
                    water_boundry_points.append(0)
                    break

                # Условие пересечения уреза с дном
                if y[i - 1] >= water_level and y[i] <= water_level:
                    x1, x2 = x[i - 1], x[i]
                    y1, y2 = y[i - 1], y[i]

                    # Нахождение координаты x уреза между точками дна
                    f = interpolate.interp1d([y1, y2], [x1, x2])
                    # Находим координату x, зная y (точка пересечения уреза с дном)
                    water_boundry_x.append(float(f(water_level)))
                    water_boundry_y.append(water_level)
                    # Присоединяем номер точки дна с границей воды
                    water_boundry_points.append(i - 1)
                    break  # Прерываем поиск если нашли пересечение

                # Условие непересекания дна и дохождения до начала участка
                elif i - 1 == 0 and y[i - 1] <= water_level:
                    water_boundry_x.append(x[i - 1])
                    water_boundry_y.append(water_level)
                    water_boundry_points.append(i - 1)
                    break  # Прерываем поиск если нашли пересечение

            # Цикл вправо от стартовой точки
            for i in range(start_point[0], len(y) - 1):
                # Условие пересечения уреза с дном
                if y[i] <= water_level and y[i + 1] >= water_level:
                    x1, x2 = x[i], x[i + 1]
                    y1, y2 = y[i], y[i + 1]

                    # Нахождение координаты x уреза между точками дна
                    f = interpolate.interp1d([y1, y2], [x1, x2])
                    # Находим координату x, зная y (точка пересечения уреза с дном)
                    water_boundry_x.append(float(f(water_level)))
                    water_boundry_y.append(water_level)
                    # Присоединяем номер точки дна с границей воды
                    water_boundry_points.append(i)
                    break  # Прерываем поиск если нашли пересечение

                elif i + 1 == len(y) - 1 and y[len(y) - 1] <= water_level:
                    water_boundry_x.append(x[len(x) - 1])
                    water_boundry_y.append(water_level)
                    water_boundry_points.append(i + 1)
                    break  # Прерываем поиск если нашли пересечение

            # Если индекс минимальной отметки совпадает с правой границой участка
            if start_point[0] == len(y) - 1 and y[start_point[0]] <= water_level:
                water_boundry_x.append(x[len(y) - 1])
                water_boundry_y.append(water_level)
                water_boundry_points.append(len(y) - 1)

            result.append([water_boundry_x, water_boundry_y,
                           water_boundry_points, 0])
        return result

    # Функция выполняющая основные вычисления по данному водному сечению
    def _calculate_parameters(self, water_boundry):
        sum_sqr = 0
        water_level = self.water_level
        x = self.x
        y = self.y
        depth = []

        # Обрабатываем урезы по две точки (со второй до третьей пропускам)
        # Вводим служебые координаты (первая и последняя точки)
        x1, x2 = water_boundry[0][0], water_boundry[0][1]
        y1, y2 = water_boundry[1][0], water_boundry[1][1]

        # Точки смоченного периметра (номера точек под урезом)
        water_section_x = x[water_boundry[2][0] + 1: water_boundry[2][1] + 1]
        water_section_y = y[water_boundry[2][0] + 1: water_boundry[2][1] + 1]

        water_section_x.insert(0, x1)
        water_section_x.insert(len(water_section_x), x2)

        water_section_y.insert(0, y1)
        water_section_y.insert(len(water_section_y), y2)

        # Если первая точка УВ выше первой точки дна, вставляем точку дна на второе место
        # TODO: Костыль для определние полигона водной поверхности для расчёта с переливом и одновременным заполнением, нужно продумать как исправить
        if config.PERELIV:  # исходные данные точек x и y по всему профилю
            if water_level > y[water_boundry[2][0]]:
                water_section_x.insert(1, x[0])
                water_section_y.insert(1, y[0])
        else:  # исходные данные точек x и y по участкам
            if water_level > y[0]:
                water_section_x.insert(1, x[0])
                water_section_y.insert(1, y[0])

        # Если последняя точка УВ выше последней точки дна, вставляем точку на предпоследнее место
        if water_boundry[3] > 1 and water_level > y[-1]:
            water_section_x.insert(len(water_section_x) - 1, x[-1])
            water_section_y.insert(len(water_section_y) - 1, y[-1])

        # Координаты x и y смоченного периметра
        self.water_section_x = water_section_x
        self.water_section_y = water_section_y

        # Определяем ширину водной поверхности
        self.width = x2 - x1

        # Площадь воды
        self.area = poly_area(water_section_x, water_section_y)

        # Глубины
        for i in range(len(water_section_y)):
            depth.append(water_level - water_section_y[i])

        # Средняя глубина
        self.average_depth = self.area / self.width # TODO: Сделать проверку деления на 0
        if self.average_depth == 0:  # Костыль
            self.average_depth = 0.00001

        # Максимальная глубина
        self.max_depth = max(depth)

        # Смоченный периметр
        for i in range(len(water_section_x) - 1):
            sum_sqr += (water_section_x[i + 1] - water_section_x[i]) ** 2
        self.w_perimeter = np.sqrt(sum_sqr)

        # Гидравлический радиус
        self.r_gidraulic = self.area / self.w_perimeter # TODO: Сделать проверку деления на 0
        if self.r_gidraulic == 0:  # Костыль
            self.r_gidraulic = 0.00001


@dataclass
class Calculation(object):
    """
    Класс гидравлических расчётов скорости, расхода воды и коэффициента Шези для водного объекта.

    :param n: Коэффициент шероховатости
    :param i: Уклон, промилле
    :param h: Средняя глубина водного сечения
    :param a: Площадь водного сечения

    """
    n: float  # Коэффициент шероховатости
    i: float  # Уклон
    h: float  # Средняя глубина
    a: float  # Площадь водного сечений
    v: float = 0  # Скорость
    q: float = 0  # Расход
    __g: float = 9.80665  # Ускорение свободного падения
    shezi: float = 0  # Коэффициент Шези
    type__: str = 'Не определен'

    def __post_init__(self):
        # В зависимости от глубины считаем по разным формулам
        # до 3-х метров по Павловскому, свыше 3-х метров по
        # Павловскому-Железнякову
        if self.h >= 0 and self.h <= 3:
            self.__shezi_pavlovskij()
        else:
            self.__shezi_pavlovskij_zheleznjakov()

        # Тип расчёта, обычная вода или селевой поток
        if config.CALC_TYPE == 1:
            # Расчёт скорости воды
            self.v = self.shezi * np.sqrt(self.h * (self.i / 1000))
        elif config.CALC_TYPE == 2:
            # Расчёт скорости воды для наносоводных селей
            self.v = 4.5 * self.h**0.67 * (self.i / 1000)**0.17
        elif config.CALC_TYPE == 3:
            # Расчёт скорости воды для грязекаменных селей селей
            self.v = 3.75 * self.h**0.50 * (self.i / 1000)**0.17
        else:
            print(
                'Ошибка выбора формулы расчёта скорости потока. Программа будет завершена.')
            sys.exit(1)
        # Расчёт расхода воды
        self.q = self.a * self.v

    # Коэффициент Шези по формуле Н. Н. Павловского, стпенной коэффициент по формуле Железнякова
    def __shezi_pavlovskij_zheleznjakov(self):
        # Показатель сетепени по формуле Г. В. Железнякова
        y = 1/np.log10(self.h) * np.log10(
            (1/2 - (self.n * np.sqrt(self.__g)/0.26) * (1 - np.log10(self.h))) +
            self.n*np.sqrt(
                1/4 * (1/self.n - np.sqrt(self.__g)/0.13 * (
                    1 - np.log10(self.h)))**2 +
                np.sqrt(self.__g)/0.13 * (
                    1/self.n + np.sqrt(self.__g) * np.log10(self.h))))

        self.shezi = (1/self.n) * self.h**y
        self.type__ = 'Коэффициент Шези определён по формуле Павловского, показатель степени определён по формуле Железнякова'

    # Коэффициент шези по формуле Маннинга
    def __shezi_mannign(self):
        self.shezi = (1/self.n) * self.h**(1/6)
        self.type__ = 'Коэффициент Шези определён по формуле Маннинга'

    # Коэффициент шези по формуле Павловского для глубин 0.1 < h < 3 (Гидрорасчёты считают по этой формуле)
    def __shezi_pavlovskij(self):
        y = 2.5 * np.sqrt(self.n) - 0.13 - 0.75 * \
            np.sqrt(self.h)*(np.sqrt(self.n) - 0.10)
        self.shezi = (1/self.n) * self.h**y
        self.type__ = 'Коэффициент шези определён по формуле Павловского для глубин 0.1 < h < 3 м'

    # Коэффициент шези по формуле Железнякова
    def __shezi_zheleznjakov(self):
        self.shezi = 1/2 * \
            (
                (1/self.n) - (np.sqrt(self.__g)/0.13) * (1 - np.log10(self.h))) + \
            np.sqrt(
                (1/4) * (1/self.n - (np.sqrt(self.__g)/0.13) * (1 - np.log10(self.h)))**2 +
                (np.sqrt(self.__g)/0.13) * ((1/self.n) +
                                            (np.sqrt(self.__g) * np.log10(self.h)))
            )
        self.type__ = 'Коэффициент шези определён по формуле Железнякова'


@dataclass
class Morfostvor(object):   

    """Класс описывающий морфствор."""
    # Основные параметры морфоствора
    title: str = ''
    x: list = field(default_factory=list)
    y: list = field(default_factory=list)
    situation: list = field(default_factory=list)
    sectors: list = field(default_factory=list)
    ele_max: float = 0
    ele_min: float = 0
    date: str = ''
    dH: int = 5
    waterline: float = 0
    erosion_limit: float = 0
    top_limit: float = 0
    top_limit_description: str = ''

    probability: list = field(default_factory=list)
    # raw_data: list = field(default_factory=list)
    coords: list = field(default_factory=list)
    strings: dict = field(default_factory=dict)

    levels_result: pd.DataFrame = pd.DataFrame
    levels_result_sectors: pd.DataFrame = pd.DataFrame
    gidraulic_result: pd.DataFrame = pd.DataFrame

    def __post_init__(self):
        # Выбор варианта расчёта
        if config.CALC_TYPE == 1:
            self.strings['type'] = 'воды'
        elif config.CALC_TYPE == 2:
            self.strings['type'] = 'наносоводного селевого потока'
        elif config.CALC_TYPE == 3:
            self.strings['type'] = 'грязекаменного селевого потока'
        else:
            print(
                'Неверно выбран тип расчёта в конфигурационном файле. Программа будет завершена.')
            sys.exit(0)

        self.qh_title = 'Кривая расхода {} Q = f(H)'.format(
            self.strings['type'])

    def read_xls(self, file_path, page=0):
        """Функция чтения из xls файла."""

        try:
            data_file = xlrd.open_workbook(file_path)  # Открываем xls файл
        except FileNotFoundError:
            print(
                'Ошибка! Файл {} не найден. Программа будет завершена.'.format(file_path))
            sys.exit(33)

        try:
            # Открываем лист по заданому номеру
            sheet = data_file.sheet_by_index(page)
        except IndexError:
            print(
                'Неверно указан индекс листа .xls файла. Проверьта параметры запуска расчёта.')
            sys.exit(1)

        sheet_title = sheet.name
        print('Считываем исходные данные из .xlsx файла: {path}, страница {page} ({title}).'.format(
            path=file_path, page=page, title=sheet_title))

        __raw_data = []  # Сырые строки xlsx файла
        i = 0

        # Позиционирование столбцов с данными в .xls файле
        __x_coord_col = 0
        __y_coord_col = 1
        __sector_name_col = 2
        __rougness_col = 3
        __slope_col = 4
        __situation_col = 5
        __descript_col = 8

        def get_sectors(self):
            """Функция считывания участоков и их параметров из исходных файлов."""

            print('    — Определяем участки ... ', end='')
            # №, Описание участка, номер первой точки, номер последней точки, коэффициент шероховатости, уклон ‰, координата x, координаты y
            lines_num = 0

            # Считываем количество строк с не пустыми координамтами
            for line in __raw_data:
                if type(line[__x_coord_col]) != str:
                    lines_num += 1

            sectors = self.sectors  # Список участков
            x = self.x  # Координаты профиля X
            y = self.y  # Координаты профиля Y

            num = 1  # Номера участокв

            ###
            # Перебираем все строки xls файла и ищем участки
            for line in range(lines_num):
                name = __raw_data[line][__sector_name_col]  # Описание профиля
                rougness = __raw_data[line][__rougness_col]  # Коэффициент шероховатости
                slope = __raw_data[line][__slope_col]  # Уклон

                # По первой строке создаём первый сектор
                if line == 0:
                    coord = []
                    sectors.append(ProfileSector(
                        num, name, line, line, rougness, slope, coord))

                # Cравниваем имя предыдущего участка с текущим, и если не совпадают то создаем новый сектор:
                elif name != sectors[num - 1].name:
                    if sectors[num - 1].id == 1:  # Если первый участок
                        # Записываем номер последний точки - 1
                        sectors[num - 1].end_point = line
                    else:  # Если все остальные участки
                        # Записываем номер последний точки в предыдщий участок для всех остальных участокв
                        sectors[num - 1].end_point = line

                    num += 1  # Увеличиваем номер сектора на 1
                    sectors.append(ProfileSector(
                        num, name, sectors[num - 2].end_point, line, rougness, slope, coord))

            # Номер последней точки в последнем секторе
            sectors[-1].end_point = len(x) - 1

            # Записываем координаты и длины участков
            for sector in sectors:
                sector.coord = (
                    (x[sector.start_point: sector.end_point + 1], y[sector.start_point: sector.end_point + 1]))  # Координаты из начальной и конечной точек
                # Длины полученнные из разницы координат по x
                sector.length = sector.get_length()

            try:
                # Максимальная отметка участка слева
                self.max_l = max(chunk_list(sector.coord[1], 2)[0])
                # Максимальная отметка участка справа
                self.max_r = max(chunk_list(sector.coord[1], 2)[1])
            except:
                print('\n\nОшибка в определении участков. Список участков:\n')
                for sector in sectors:
                    print(sector)

                print('Завершаем программу.')
                sys.exit(3)

            print('успешно, найдено {num} участка.\n'.format(num=len(sectors)))
            return sectors

        # Перебираем все строки
        # И получаем список сырых данных
        for rownum in range(1, sheet.nrows):
            row = sheet.row_values(rownum)
            __raw_data.append(row)  # Записываем данные

        # Устанавливаем основные параметры морфоствора
        print('    — Устанавливаем основные параметры морфоствора ... ', end='')
        self.title = __raw_data[2][__descript_col]  # Заголовк профиля
        self.date = __raw_data[3][__descript_col]  # Дата профиля

        self.waterline = __raw_data[4][__descript_col]  # Отметка уреза воды
        # Проверяем задан ли урез текстом, если нет округляем до 2 знаков        
        if type(self.waterline) is not str:
            self.waterline = round(self.waterline, 2)

        
        self.dH = __raw_data[5][__descript_col]  # Расчётный шаг по глубине
        self.coords = __raw_data[6][__descript_col]  # Координаты
        self.erosion_limit = __raw_data[7][__descript_col]  # Предел размыва
        self.top_limit = __raw_data[8][__descript_col]  # Верхняя граница
        # Описание верхней границы
        self.top_limit_description = __raw_data[9][__descript_col]
        print('успешно!')

        # Считываем и записываем все точки x и y профиля
        print('    — Считываем координаты профиля ... ', end='')
        for i in range(len(__raw_data)):
            if type(__raw_data[i][__x_coord_col]) != str:
                self.x.append(__raw_data[i][__x_coord_col])
                self.y.append(__raw_data[i][__y_coord_col])
                self.situation.append(__raw_data[i][__situation_col])
        print('успешно, найдено {} точки, длина профиля {} м'.format(
            len(self.x), self.x[-1]))

        self.ele_min = min(self.y)  # Минимальная отметка профиля
        self.ele_max = max(self.y)  # Максимальная отметка профиля

        # Заполнения таблицы обеспеченностей
        print('    — Считываем обеспеченности ... ', end='')
        for i in range(6, len(__raw_data[0])):
            self.probability.append([__raw_data[0][i], __raw_data[1][i]])

        # Удаляем пустые обеспеченности из списка обеспеченностей
        self.probability = [x for x in self.probability if x != ['', '']]

        print('успешно, найдено {} обспеченностей.'.format(len(self.probability)))

        # Обработка и получение данных по секторам из "сырых" данных
        self.sectors = get_sectors(self)

    def get_min_sector(self):
        """
        Функция нахождения участка с наименьшей отметкой дна.

        :return: [Номер по списку, [Участок]]
        """

        id = 0
        i = 0
        min_sector = self.sectors[0]

        for sector in self.sectors:
            if min(sector.coord[1]) < min(min_sector.coord[1]):
                min_sector = sector
                id = i
            i = i + 1
        return [id, min_sector]

    def get_q_max(self):
        """
        Функция нахождения максимальной обеспеченности и расхода воды по исходным данным.

        :return: [Обеспеченность, Расход]
        """
        q_max = float(self.probability[0][1])
        obsp = self.probability[0][0]
        for Q in self.probability:
            if q_max <= Q[1]:
                q_max = Q[1]
                obsp = Q[0]

        return [obsp, q_max]

    def doc_export(self, out_filename, r=False):
        print('\n\nФормируем doc файл: ')
        doc_file = out_filename

        if r:
            doc = DocxTemplate('assets/report_template.docx')
        else:
            if os.path.isfile(doc_file):
                doc = Document(doc_file)
                paragraphs = doc.paragraphs
                run = paragraphs[-1].add_run()
                run.add_break(WD_BREAK.PAGE)
            else:
                if config.REWRITE_DOC_FILE:
                    print(
                        '    — Включена перезапись файла, удаляем старый и создаём новый.')
                else:
                    print('    — Файл не найден! Создаём новый.')
                doc = DocxTemplate('assets/report_template.docx')

        if config.GIDRAULIC_CURVE:
            self.fig_QH = GraphQH(self)
        if config.SPEED_CURVE:
            self.fig_QV = GraphQV(self)

        self.fig_profile = GraphProfile(self)

        # Отрисовка смоченного периметра
        if config.PROFILE_WET_PERIMITER:
            self.fig_profile.draw_wet_perimeter()

        # Отрисовка верхней границы сооружения
        if self.top_limit:        
            self.fig_profile.draw_top_limit(self.top_limit, text=self.top_limit_description)

        # Отрисовка границы предельного размыва профиля
        if self.erosion_limit:
            self.fig_profile.draw_erosion_limit(self.erosion_limit)

        # Отрисовка уровней воды на графике профиля
        self.fig_profile.draw_levels_on_profile(self.levels_result)


        self.fig_profile._update_limit()
        if self.waterline and type(self.waterline) != str:
            self.fig_profile.draw_waterline(round(self.waterline, 2), color='blue', linestyle='-')

        # Проверяем наличие временной папки
        try:
            os.stat('temp')
        except:
            os.mkdir('temp')

        print('    — Сохраняем график профиля ... ', end='')
        self.fig_profile.fig.savefig('temp/Profile.png', dpi=config.FIG_DPI)
        print('успешно!')

        # Вставляем заголовк профиля
        doc.add_paragraph(self.title, style='З-приложение-подзаголовок')

        # Добавляем изображения профиля и гидравлической кривой
        print('    — Вставляем графику (профиль и кривую)... ', end='')
        doc.add_picture('temp/Profile.png', width=Cm(16.5))
        setLastParagraphStyle('Р-рисунок', doc)

        if config.GRAPHICS_TITLES_TEXT:
                doc.add_paragraph('Рисунок — ' + self.fig_profile.morfostvor.title, style='Р-название')

        if config.GIDRAULIC_CURVE:
            print('    — Сохраняем график гидравлической кривой ... ', end='')
            self.fig_QH.fig.savefig('temp/QH.png', dpi=config.FIG_DPI)
            print('успешно!')

            doc.add_picture('temp/QH.png', width=Cm(16.5))
            setLastParagraphStyle('Р-рисунок', doc)

            if config.GRAPHICS_TITLES_TEXT:
                doc.add_paragraph('Рисунок — ' + self.fig_QH._ax_title_text, style='Р-название')

        if config.SPEED_CURVE:
            print('    — Сохраняем график кривой скоростей ... ', end='')
            self.fig_QV.fig.savefig('temp/QV.png', dpi=config.FIG_DPI)
            print('успешно!')

            doc.add_picture('temp/QV.png', width=Cm(16.5))
            setLastParagraphStyle('Р-рисунок', doc)
            print('успешно!')

            if config.GRAPHICS_TITLES_TEXT:
                doc.add_paragraph('Рисунок — ' + self.fig_QV._ax_title_text, style='Р-название')


        # Вывод таблицы расчётных уровней воды
        print('    — Записываем таблицу уровней воды ... ', end='')
        param = (('Обеспеченность P, %', 'Расход Q, м³/сек', 'Уровень H, м'),
                 (5.6, 5.6, 5.6),
                 (':g', ':g', ':.2f'))
        print('успешно!')

        # levels_result = self.levels_result
        levels_result = self.levels_result[[
            'P', 'Q', 'H']].round(3).values.tolist()
        write_table(doc, levels_result, param,
                    'Таблица - Расчётные уровни {}'.format(self.strings['type']))

        # Вывод таблицы участков
        print('    — Записываем таблицу участков ... ', end='')
        param = (('№', 'Описание', 'Уклон i, ‰', 'Коэффициент шероховатости n'),
                 (1.5, 5.1, 5.1, 5.1),
                 (':d', '', ':g', ':.3f'))
        sectors = []

        i = 1
        for sector in self.sectors:
            sectors.append([i, sector.name, sector.slope, sector.roughness])
            i += 1

        write_table(doc, sectors, param,
                    'Таблица - Расчётные участки и их параметры')
        print('успешно!')

        print('    — Записываем таблицу кривой расхода воды ... ', end='')

        # Вывод таблицы гидравлической кривой
        param = (('Отм. уровня H, м', 'Площадь F, м²', 'Ширина B, м', 'Средняя глубина Hср, м', 'Макс. глубина Hмакс, м',
                  'Средняя скорость Vср, м/сек', 'Расход Q, м³/сек'),
                 (5, 5, 5, 5, 5, 5, 5),
                 (':.2f', ':.3f', ':.3f', ':.3f', ':.3f', ':.3f', ':.3f'))

        table = self.gidraulic_result
        table_round = table.round(3)  # Округляем
        # Убираем столбец с коэффициентами Шези
        table_round = table_round.drop(columns=['Shezi'])

        if config.DOC_TABLE_SHORT:
            # Количество строк в таблице
            table_quant = table_round['УВ'].count()

            # Уменьшаем количество выводимых строк в таблицу
            # чтобы поместилось на один лист
            if table_quant <= 25:
                divider = 1
            elif table_quant > 25 and table_quant <= 50:
                divider = 2
            elif table_quant > 50 and table_quant <= 75:
                divider = 3
            elif table_quant > 75 and table_quant <= 100:
                divider = 4
            elif table_quant > 100 and table_quant <= 125:
                divider = 5
            else:
                divider = 10
        else:
            divider = 1

        # Записываем только чётные элементы таблицы
        table_round = table_round[table_round.index % divider == 0]
        summ_gidraulic = table_round.values.tolist()  # Переводим в список
        write_table(doc, summ_gidraulic, param,
                    'Таблица - Параметры расчёта кривой расхода {}'.format(self.strings['type']))
        doc.add_paragraph('Расчётный шаг: {:g} см. В таблице приведён каждый {}-й результат расчёта.'.format(
            self.dH, divider), style='Т-примечание')
        print('успешно!')

        # Удаляем объект профиля
        self.fig_profile.clean()

        try:
            doc.save(doc_file)
        except PermissionError:
            print(
                '\nОшибка! Не удалось сохранить файл. Проверьте возможность записи файла по указанному пути.')
            print('Возможно записываемый файл уже существует и открыт.')
            sys.exit(1)

        try:
            os.remove('temp/QH.png')
            os.remove('temp/QV.png')
            os.remove('temp/Profile.png')
            os.rmdir('temp')
        except:
            pass

        print(
            '\nФайл {filename} сохранён успешно.\n-------------------------------------\n'.format(filename=doc_file))

    def calculate(self):
        # Значение расхода до которого необходимо считать (максимальной введенная обеспеченности + 20%)
        consumption_check = self.get_q_max()[1] + (
            self.get_q_max()[1] * 0.20)

        # Проверяем задан ли расчётный шаг в исходных данных
        if isinstance(self.dH, str) or self.dH == 0:
            self.dH = 1
            dH = self.dH
        else:
            dH = self.dH

        # Переводим сантиметры приращения в метры
        dH = dH / 100

        min_sector = self.get_min_sector()

        # Исходные сектора для расчёта (сектор, содержащий минимальную отметку)
        calc_sectors = [min_sector[0]]

        # Уровень воды, с минимальным отступом
        water_level = min(self.y) + dH

        # Обнулённые переменные
        consumption_summ = 0
        area_summ = 0
        n = 0
        result = {}
        summ = []
        first_calc = True

        # Первый расчётный элемент суммирующей кривой со всеми нулями
        summ.append([
            min(self.y),
            0,
            0,
            0,
            0,
            0,
            0,
            0
        ])

        # Цикл расчёта до максимальной обеспеченности + 20% из исходных данных
        while consumption_summ < consumption_check:
            print('Выполняем расчёты для уровня {water_level:.2f}'.format(
                water_level=water_level), end='\r')

            consumption_summ = 0
            wc_list = list()
            area_list = list()
            summ_result = [[], [], [], [], [], [], [], []]

            if config.PERELIV:
                for i in calc_sectors:
                    sector = self.sectors[i]
                    x = sector.coord[0]
                    y = sector.coord[1]

                    # Максимальная отметка слева
                    previus_min_ele = max(chunk_list(y, 2)[0])
                    # Максимальная отметка справа
                    next_min_ele = max(chunk_list(y, 2)[1])

                    # Проверка на перелив через границы участка
                    if (water_level >= previus_min_ele) and (i - 1 not in calc_sectors) and (i - 1 >= 0):
                        calc_sectors.append(i - 1)
                    if (water_level >= next_min_ele) and (i + 1 not in calc_sectors) and (i + 1 <= len(self.sectors) - 1):
                        calc_sectors.append(i + 1)

                    # Сектор воды и основные его параметры
                    # Расчетный участок является участком с минимальными отметками
                    # либо расчёт выполняется с одновременным заполнением
                    # начинаем заполнять с точки с минимальной отметкой
                    if sector.id == min_sector[1].id:
                        water = WaterSection(x, y, water_level)

                    # Расчетный участок находится слева от начального
                    # начинаем заполнять с крайней правой точки
                    elif sector.id < min_sector[1].id:
                        water = WaterSection(x, y, water_level, start_point=[
                                             len(y) - 1, y[-1]])

                    # Расчетный участок находится справа от начального
                    # начинаем заполнять с крайней левой точки
                    elif sector.id > min_sector[1].id:
                        water = WaterSection(
                            x, y, water_level, start_point=[0, y[0]])

                    # Расчёт параметров для воды
                    cc = Calculation(
                        h=water.average_depth, n=sector.roughness, i=sector.slope, a=water.area)

                    # Добавляем в список с результирующими значениями значения по секторам для последующего суммирования/вычисления средних значений
                    # TODO: создать датафрейм с результатами
                    summ_curve = pd.DataFrame(
                        data=summ,
                        columns=['УВ', 'F', 'B', 'Hср',
                                 'Hмакс', 'V', 'Q', 'Shezi']
                    )

                    summ_result[0] = water_level  # H
                    summ_result[1].append(water.area)  # F
                    summ_result[2].append(water.width)  # B
                    summ_result[3].append(water.average_depth)  # Hср
                    summ_result[4].append(water.max_depth)  # Hмакс
                    summ_result[5].append(cc.v)  # V
                    summ_result[6].append(cc.q)  # Q
                    summ_result[7].append('{sector}: {shezi}'.format(
                        sector=sector.name, shezi=cc.shezi))  # Шези

                    wc_list.append(cc.q)

                    # Если это первый расчёт, записываем нулевые значения
                    if first_calc:
                        result[sector.name] = list()
                        result[sector.name].append(
                            [min(self.y), 0, 0, 0, 0, 0, 0])
                        first_calc = False

                    # Записываем значения по каждому сектору в отдельный список
                    try:
                        result[sector.name].append(
                            [water_level, water.area, water.width, water.average_depth, water.max_depth, cc.v,
                             cc.q, cc.v])
                    except KeyError:
                        result[sector.name] = list()
                        result[sector.name].append(
                            [water_level, water.area, water.width, water.average_depth, water.max_depth, cc.v,
                             cc.q, cc.shezi])

            else:
                # Расчёт с заполнением по участкам
                for sector in self.sectors:
                    x = sector.coord[0]
                    y = sector.coord[1]

                    if min(y) < water_level:
                        # Сектор воды и основные его параметры
                        water = WaterSection(x, y, water_level)

                        # Расчёт параметров для воды
                        cc = Calculation(
                            h=water.average_depth, n=sector.roughness, i=sector.slope, a=water.area)

                        # Добавляем в список с результирующими значениями значения по секторам для последующего суммирования/вычисления средних значений
                        summ_result[0] = water_level  # H
                        summ_result[1].append(water.area)  # F
                        summ_result[2].append(water.width)  # B
                        summ_result[3].append(water.average_depth)  # Hср
                        summ_result[4].append(water.max_depth)  # Hмакс
                        summ_result[5].append(cc.v)  # V
                        summ_result[6].append(cc.q)  # Q
                        summ_result[7].append('{sector}: {shezi}'.format(
                            sector=sector.name, shezi=cc.shezi))  # Шези

                        wc_list.append(cc.q)

                        # Если это первый расчёт, записываем нулевые значения
                        if first_calc:
                            result[sector.name] = list()
                            result[sector.name].append(
                                [min(self.y), 0, 0, 0, 0, 0, 0])
                            first_calc = False

                        # Записываем значения по каждому сектору в отдельный список
                        try:
                            result[sector.name].append(
                                [water_level, water.area, water.width, water.average_depth, water.max_depth, cc.v,
                                 cc.q, cc.v])
                        except KeyError:
                            result[sector.name] = list()
                            result[sector.name].append(
                                [water_level, water.area, water.width, water.average_depth, water.max_depth, cc.v,
                                 cc.q, cc.shezi])

            consumption_summ += sum(wc_list)
            area_summ += sum(area_list)

            # Вычисляем параметры суммирующей кривой
            wl = summ_result[0]  # Уровень воды
            F = sum(summ_result[1])  # Площадь водного сечения
            B = sum(summ_result[2])  # Ширина русла
            Hsr = F / B  # Средняя глубина
            Hmax = max(summ_result[4])  # Макс глубина
            Q = sum(summ_result[6])  # Расход
            Vsr = Q / F  # Средняя Скорость
            Sh = summ_result[7]  # Коэффициент шези

            summ.append([
                wl,
                F,
                B,
                Hsr,
                Hmax,
                Vsr,
                Q,
                Sh
            ])

            water_level += dH
            n += 1

        df = pd.DataFrame(
            data=summ,
            columns=['УВ', 'F', 'B', 'Hср', 'Hмакс', 'V', 'Q', 'Shezi']
        )

        df_list = {}

        for sector in result:
            df_list[sector] = pd.DataFrame(
                data=result[sector],
                columns=['УВ', 'F', 'B', 'Hср', 'Hмакс', 'V', 'Q', 'Shezi']
            )

        # Находим H от Q
        result = pd.DataFrame(
            columns=['P', 'Q', 'H']
        )

        for element in self.probability:
            fQ = interpolate.interp1d(df['Q'], df['УВ'])
            fV = interpolate.interp1d(df['Q'], df['V'])
            h = float(fQ(element[1]))
            v = float(fV(element[1]))

            result = result.append(
                {'P': element[0],
                 'H': h,
                 'Q': element[1],
                 'V': v,
                 }, ignore_index=True
            )
        self.levels_result = result
        self.levels_result_sectors = df_list
        self.gidraulic_result = df


@dataclass
class Graph(object):
    _fig_size = (16.5, 11)
    _y_limits = []
    _fig_num = 0
    _y_lim = (0, 100)

    _x_label_text = ''
    _y_label_text = ''
    _ax_title_text = ''

    morfostvor: Morfostvor = Morfostvor
    fig: plt.figure = plt.figure(_fig_num, figsize=_fig_size)
    ax: plt.subplot = fig.add_subplot(111) 

    def __post_init__(self):
        self.clean()
        morfostvor = self.morfostvor

        # Вытягиваем цвета
        self.sector_colors = {}
        for sector in morfostvor.sectors:
            self.sector_colors[sector.name] = sector.color
        
        # Выполняем отрисовку содержимого
        self.draw()
        self.set_style()
    
    def draw(self):
        pass

    def set_style(self):
        fig = self.fig
        ax = self.ax

        fig.subplots_adjust(bottom=0.08, left=0.08, right=0.9)

        # Устанавливаем заголовки графиков
        if config.GRAPHICS_TITLES:
            ax.set_title(
                self._ax_title_text, color=config.COLOR['title_text'], fontsize=config.FONT_SIZE['title'], y=1.1)

        # Настравиааем границы и толщину линий границ
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_linewidth(config.LINE_WIDTH['ax_border'])
        ax.spines['bottom'].set_linewidth(config.LINE_WIDTH['ax_border'])

        # Устанавливаем отступы в графиках
        ax.margins(0.025)

        # Устанавливаем параметры засечек на основых осях
        ax.tick_params(
            which='major',
            direction='out',
            width=2,
            length=5,
            pad=10,
            labelcolor=config.COLOR['ax_value_text'],
            labelsize=config.FONT_SIZE['ax_major'])

        ax.tick_params(
            which='minor',
            direction='out',
            width=1.5,
            length=3.5,
            pad=10,
            labelcolor=config.COLOR['ax_value_text'],
            labelsize=config.FONT_SIZE['ax_minor'])

        # Устанавливаем параметры подписей осей
        ax.set_xlabel(self._x_label_text, color=config.COLOR['ax_label_text'],
                        fontsize=config.FONT_SIZE['ax_label'], fontstyle='italic')
        ax.xaxis.set_label_coords(1.05, -0.025)
        ax.set_ylabel(self._y_label_text, color=config.COLOR['ax_label_text'],
                        fontsize=config.FONT_SIZE['ax_label'], fontstyle='italic')
        ax.yaxis.set_label_coords(-0.025, 1.08)

        # Устанавливает параметры вывода значений осей
        ax.yaxis.set_major_formatter(
            matplotlib.ticker.FormatStrFormatter('%.10g'))

        # Настройка параметров отображение сетки
        ax.grid(which='major',
                color=config.COLOR['ax_grid'], linestyle=':', linewidth=1, alpha=0.9)
        ax.grid(which='minor',
                color=config.COLOR['ax_grid_sub'], linestyle=':', linewidth=1, alpha=0.9)

    def clean(self):
        """Очистка осей графика и обнуление связанных переменых
        """
        # Очищаем все оси
        for ax in vars(self):
            if ax.startswith('ax'):
                command = "self." + ax + ".cla()"
                exec(command)
                
        # Обнуляем границы y
        self._y_limits = []
        self._y_limits = []


@dataclass
class GraphQH(Graph):
    # Номер рисунка
    _fig_num = 2

    # Подписи осей
    _x_label_text = 'Q, м³/с'
    _y_label_text = 'H, м'
    _ax_title_text = 'Гидравлическая кривая'

    def draw_gidraulic_curve(self):
        morfostvor = self.morfostvor
        ax = self.ax
        result_sectors = morfostvor.levels_result_sectors

        # Отрисовка суммирующей кривой на графике
        ax.plot(morfostvor.gidraulic_result['Q'],
                morfostvor.gidraulic_result['УВ'], label='Сумма', linewidth=3, color='red')

        # Отрисовка кривых по участкам
        for sector in result_sectors:
            ax.plot(result_sectors[sector]['Q'], result_sectors[sector]['УВ'], '--',
                    label=sector, color=self.sector_colors[sector])  # marker='o', markersize='3',

        ax.legend(loc='lower right', fontsize=config.FONT_SIZE['legend'])

    def draw_water_levels(self):
        morfostvor = self.morfostvor
        ax = self.ax
        # Вывод расчётных уровней
        try:
            if config.GIDRAULIC_CURVE_LEVELS:
                min_h = min(morfostvor.gidraulic_result['УВ'])

                for index, row in morfostvor.levels_result.iterrows():
                    x1, x2 = 0, row['Q']
                    y1, y2 = row['H'], row['H']

                    # Вывод значений округленных, проверка на содержание значений
                    try:
                        water_level_text = ax.text(0.002, row['H'], '▼$P_{{{:.2g}\%}} = {H:.2f}$'.format(
                            row['P'], H=row['H']), color=config.COLOR['water_level_text'], fontsize=config.FONT_SIZE['water_level'], weight='bold')
                        water_level_text.set_path_effects([path_effects.Stroke(
                            linewidth=3, foreground='white', alpha=0.55), path_effects.Normal()])
                    except ValueError:
                        water_level_text = ax.text(0.002, row['H'], '▼{} = {H:.2f}'.format(
                            row['P'], H=row['H']), color=config.COLOR['water_level_text'], fontsize=config.FONT_SIZE['water_level'], weight='bold')
                        water_level_text.set_path_effects([path_effects.Stroke(
                            linewidth=3, foreground='white', alpha=0.55), path_effects.Normal()])

                    ax.plot([x1, x2, x2, x2], [y1, y2, min_h, min_h], linestyle='-',
                            color='mediumturquoise', marker='o', linewidth=1, markersize=1)
        except:
            print('Внимание! Вывод расчётных уровней на график не возможен!')
        pass

    def draw(self):
        self.draw_gidraulic_curve()
        self.draw_water_levels()        


@dataclass
class GraphQV(Graph):
    # Номер рисунка
    _fig_num = 3
    _fig_size = (16.5, 11)
    fig: plt.figure = plt.figure(_fig_num, figsize=_fig_size)
    ax: plt.subplot = fig.add_subplot(111)

    # Подписи осей
    _x_label_text = 'Q, м³/с'
    _y_label_text = 'V, м/c'
    _ax_title_text = 'Кривая скоростей'

    def draw_speed_curve(self):
        morfostvor = self.morfostvor
        ax = self.ax
        result_sectors = morfostvor.levels_result_sectors

        # Отрисовка суммирующей кривой на графике
        ax.plot(morfostvor.gidraulic_result['Q'],
                morfostvor.gidraulic_result['V'], label='Сумма', linewidth=3, color='red')
        
        # Отрисовка кривых по участкам
        for sector in result_sectors:
            ax.plot(result_sectors[sector]['Q'], result_sectors[sector]['V'], '--',
                    label=sector, color=self.sector_colors[sector])  # marker='o', markersize='3',
        
        ax.legend(loc='lower right', fontsize=config.FONT_SIZE['legend'])

    def draw(self):
        self.draw_speed_curve()


@dataclass
class GraphProfile(Graph):
    _fig_size =(16.5, 12)
    _fig_num = 1

    fig: plt.figure = plt.figure(_fig_num, figsize=_fig_size)

    __gs = gridspec.GridSpec(80, 3)

    ax_top: plt.subplot = fig.add_subplot(__gs[0, :], frame_on=False)
    ax: plt.subplot = fig.add_subplot(__gs[1:62, :])
    ax_bottom: plt.subplot = fig.add_subplot(__gs[62:, :])
    ax_bottom_overlay: plt.subplot = fig.add_subplot(
        __gs[62:, :], frame_on=False)   


    def __post_init__(self):
        self.clean()

        # Настройка параметров графиков и их инициализация
        self.fig.subplots_adjust(bottom=0.08, left=0.08, right=0.9)

        # Добавляем в список границ максимальную и минимальную отметки
        self._y_limits.append(max(self.morfostvor.y))
        self._y_limits.append(min(self.morfostvor.y))

        self._update_limit()
        self.set_style()

        self.draw_profile_footer()
        self.draw_sectors()
        self.draw_profile_bottom()

    def draw_profile_bottom(self):
        """
        Отрисовка дна профиля.

        :return: Отрисовыает дно профиля на графике ax_profile.
        """

        self.ax.plot(self.morfostvor.x, self.morfostvor.y,
                        color=config.COLOR['profile_bottom'], linewidth=config.LINE_WIDTH['profile_bottom'], linestyle='solid', )

    def draw_profile_footer(self):
        """
        Отрисовка подвала с информацией о профиле.

            :param self: 
        """
        # Подписи Данных в подвале
        if config.PROFILE_LEVELS_TABLE:
            self.ax_bottom.set_xlabel(
                'ПК, м\n\nОтм. м\n\nРасст. м\n\nКоэфф. n',
                color=config.COLOR['bottom_text'],
                fontsize=config.FONT_SIZE['bottom_description'],
                fontstyle='italic', horizontalalignment='left')

            self.ax_bottom.xaxis.set_label_coords(1.02, 0.92)
        # Горизонтальные разделители подвала (полная рамка)
        self.ax_bottom_overlay.plot((self.morfostvor.x[0], self.morfostvor.x[-1]), (0, 0), color=config.COLOR['border'],
                                    linewidth=config.LINE_WIDTH['profile_bottom'], linestyle='solid')
        self.ax_bottom_overlay.plot((self.morfostvor.x[0], self.morfostvor.x[-1]), (5, 5), color=config.COLOR['border'],
                                    linewidth=config.LINE_WIDTH['profile_bottom'], linestyle='solid')
        self.ax_bottom_overlay.plot((self.morfostvor.x[0], self.morfostvor.x[-1]), (10, 10),
                                    color=config.COLOR['border'], linewidth=config.LINE_WIDTH['profile_bottom'], linestyle='solid')
        self.ax_bottom_overlay.plot((self.morfostvor.x[0], self.morfostvor.x[-1]), (15, 15),
                                    color=config.COLOR['border'], linewidth=config.LINE_WIDTH['profile_bottom'], linestyle='solid')
        self.ax_bottom_overlay.plot(
            (self.morfostvor.x[0], self.morfostvor.x[-1]), (20, 20), alpha=0)

        # Технический разделитель (для увеличения размера границ)
        self.ax_bottom.plot(
            (self.morfostvor.x[0], self.morfostvor.x[0]), (30, 40), alpha=0)

        # Цикл по всем точкам
        for i in range(len(self.morfostvor.x)):
            x = self.morfostvor.x[i]
            y = self.morfostvor.y[i]

            # Разделители расстояний между точками
            self.ax_bottom.plot(
                (x, x), (10, 20), color=config.COLOR['border'], linewidth=config.LINE_WIDTH['profile_bottom'], linestyle='solid')

            # Подписи отметок
            self.ax_bottom.text(x, 25, '{:.2f}'.format(
                y), color=config.COLOR['bottom_text'], fontsize=config.FONT_SIZE['bottom_small'], verticalalignment='center', horizontalalignment='center', rotation='90')

        # Цикл по точкам до предпоследней
        for i in range(len(self.morfostvor.x) - 1):
            x = self.morfostvor.x[i]
            x1 = self.morfostvor.x[i + 1]
            y = self.morfostvor.y[i]

            # Подписи расстояний между точками
            self.ax_bottom.text((x + x1) / 2, 15, '{:d}'.format(round(
                x1 - x)), color=config.COLOR['bottom_text'], fontsize=config.FONT_SIZE['bottom_main'], verticalalignment='center', horizontalalignment='center',)

        # Цикл по участкам
        for sector in self.morfostvor.sectors:
            x = self.morfostvor.x[sector.start_point]
            x1 = self.morfostvor.x[sector.end_point]

            # Подписи коэффициентов шероховатости по участкам
            try:
                self.ax_bottom.text((x + x1) / 2, 5, '{:.3f}'.format(
                    sector.roughness),  color=config.COLOR['bottom_text'], fontsize=config.FONT_SIZE['bottom_main'], verticalalignment='center', horizontalalignment='center')
            except ValueError:
                print('\nОшибка в указании параметров участков (коэффициент шероховатости или разделение на участки). Проверить данные.')
                sys.exit(1)

            # Разделители коэффициентов шероховатости
            self.ax_bottom.plot(
                (x, x), (0, 10), color=config.COLOR['border'], linewidth=config.LINE_WIDTH['profile_bottom'], linestyle='solid')
            self.ax_bottom.plot(
                (x1, x1), (0, 10), color=config.COLOR['border'], linewidth=config.LINE_WIDTH['profile_bottom'], linestyle='solid')

    def draw_sectors(self):
        """
        Отрисовка различной информации связанной с участками профиля.

        :param fill: [bool] - заливка полигонов участков на профиле соответствующими цветами
        :param bottom: [bool] - заливка линии дна соответствующими участкам цветами
        :param label: [bool] - отрисовка названий участков, их длин и стрелок обозначающих границы участков
        :return: Отрисовка графической информации по участкам профиля на графике ax_profile.
        """

        h_max = np.floor(max(self.morfostvor.y)) + 1

        for sector in self.morfostvor.sectors:
            points = []

            for i in range(len(sector.coord[0])):
                points.append((sector.coord[0][i], sector.coord[1][i]))

            points.insert(0, (sector.coord[0][0], h_max))
            points.append((sector.coord[0][-1], h_max))

            polygon = matplotlib.patches.Polygon(
                points, alpha=0.04, linestyle='--', label=sector.name)
            polygon.set_color(sector.color)

            # Подписи названий и длин участокв со стрелками
            if config.PROFILE_SECTOR_LABEL:
                p0 = 1
                p1 = 2
                p3 = 3

                # Расчёт середины участка (для центровки текста)
                cent_x = sector.coord[0][-1] - \
                    ((sector.coord[0][-1] - sector.coord[0][0]) / 2)

                # Вывод ширины участка
                self.ax_top.text(
                    cent_x, p1, '{:d} м'.format(round(sector.length)), color=config.COLOR['sector_text'],
                    verticalalignment='center', horizontalalignment='center',
                    bbox={'facecolor': 'white', 'edgecolor': 'white', 'alpha': 1, 'pad': 2.5})

                self.ax_top.text(cent_x, 6, sector.name, color=config.COLOR['sector_text'],
                                    verticalalignment='center', horizontalalignment='center',)

                # Вывод разделителя участков профиля
                self.ax_top.plot([sector.coord[0][0], sector.coord[0][0]], [p0, p3], color=config.COLOR['sector_line'],
                                    linestyle='-', linewidth=config.LINE_WIDTH['sector_line'])  # Горизонтальная слева
                self.ax_top.plot([sector.coord[0][-1], sector.coord[0][-1]], [p0, p3], color=config.COLOR['sector_line'],
                                    linestyle='-', linewidth=config.LINE_WIDTH['sector_line'])  # Горизонтальная справа

                self.ax_top.plot([sector.coord[0][0], cent_x], [p1, p1], color=config.COLOR['sector_line'],
                                    linestyle='-', linewidth=config.LINE_WIDTH['sector_line'])  # Вертикальная слева
                self.ax_top.plot([cent_x, sector.coord[0][-1]], [p1, p1], color=config.COLOR['sector_line'],
                                    linestyle='-', linewidth=config.LINE_WIDTH['sector_line'])  # Вертикальная справа

            # Заливка на профиле участков
            if config.PROFILE_SECTOR_FILL:
                self.ax.add_patch(polygon)

            # Цвет линии дна по участкам
            if config.PROFILE_SECTOR_BOTTOM_LINE:
                self.ax.plot(
                    sector.coord[0], sector.coord[1], '-', color=sector.color)

    def set_style(self):
        # Устанавливаем заголовки графиков
        if config.GRAPHICS_TITLES:        
            self.ax.set_title(
                self.morfostvor.title, color=config.COLOR['title_text'], fontsize=config.FONT_SIZE['title'], y=1.1)

        self.ax.set_ylim(self._y_lim)


        # Настравиааем границы и толщину линий границ
        self.ax.spines['top'].set_visible(False)
        self.ax.spines['right'].set_visible(False)
        self.ax.spines['left'].set_linewidth(config.LINE_WIDTH['ax_border'])
        self.ax.spines['bottom'].set_linewidth(config.LINE_WIDTH['ax_border'])

        self.ax_bottom.spines['top'].set_visible(False)
        self.ax_bottom.spines['right'].set_linewidth(
            config.LINE_WIDTH['ax_border'])
        self.ax_bottom.spines['left'].set_linewidth(
            config.LINE_WIDTH['ax_border'])
        self.ax_bottom.spines['bottom'].set_linewidth(
            config.LINE_WIDTH['ax_border'])

        # Устанавливаем отступы в графиках
        self.ax.margins(0.025)
        self.ax_top.margins(0.025)
        self.ax_bottom.margins(0.025, 0)
        self.ax_bottom_overlay.margins(0)

        # Устанавливаем прозрачность заливки фона
        self.ax_top.patch.set_alpha(0)
        self.ax_bottom.patch.set_alpha(0)
        self.ax_bottom_overlay.patch.set_alpha(0)

        # Включаем отображение сетки
        self.ax.grid(True, which='both')

        # Включаем отображение второстепенных засечек на осях
        self.ax.minorticks_on()

        # Устанавливаем параметры засечек на основых осях
        self.ax.tick_params(
            which='major',
            direction='out',
            width=2,
            length=5,
            pad=13,
            labelcolor=config.COLOR['ax_label_text'],
            labelsize=config.FONT_SIZE['ax_major'])

        self.ax.tick_params(
            which='minor',
            direction='out',
            width=1.5,
            length=3,
            pad=10,
            labelcolor=config.COLOR['ax_label_text'],
            labelsize=config.FONT_SIZE['ax_minor'])

        # Отключаем засечки и подписи на осях вспомогательных графиков
        self.ax_bottom.set_xticks([])
        self.ax_bottom.set_yticks([])
        self.ax_bottom_overlay.set_xticks([])
        self.ax_bottom_overlay.set_yticks([])
        self.ax_top.set_xticks([])
        self.ax_top.set_yticks([])

        # Устанавливаем параметры подписей осей
        self.ax.set_ylabel(
            'H, м', color=config.COLOR['ax_label_text'], fontsize=config.FONT_SIZE['ax_label'], fontstyle='italic')
        self.ax.yaxis.set_label_coords(-0.025, 1.08)

        # Устанавливает параметры вывода значений осей
        self.ax.yaxis.set_major_formatter(
            matplotlib.ticker.FormatStrFormatter('%.10g'))

        # Настройка параметров отображение сетки
        self.ax.grid(
            which='major', color=config.COLOR['ax_grid'], linestyle=':', linewidth=1, alpha=0.9)
        self.ax.grid(
            which='minor', color=config.COLOR['ax_grid_sub'], linestyle=':', linewidth=1, alpha=0.9)


    def draw_profile_point_lines(self):
        """
        Отрисовка вертикальных линий от точек до подвала.

        """
        for i in range(len(self.morfostvor.x)):
            self.ax.plot(
                (self.morfostvor.x[i], self.morfostvor.x[i]),
                (self.morfostvor.y[i], self._y_lim[0]),
                color=config.COLOR['profile_point_line'],
                linewidth=config.LINE_WIDTH['profile_point_line'],
                linestyle='solid')

    def draw_erosion_limit(self, h, x1=None, x2=None, text='▼$H_{{разм.}} = {h:.2f}$'):
        """Функция отрисовки линии предельного размыва профиля.
        
        Arguments:
            h {[float]} -- Отметка линии предельного размыва
        
        Keyword Arguments:
            x1 {[float]} -- Координата начала линии (default: {None})
            x2 {[float]} -- Координа конца линии (default: {None})
            text {[string]} -- Текст подписи линии (default: {'▼$H_{{разм.}} = {h:.2f}$'})
        """
        # if config.PROFILE_EROSION_LIMIT and not isinstance(self.morfostvor.erosion_limit, str):
        # Исходное ограничение линии предельного размыва по всему профилю

        # Если координаты начала и конца линии не заданы, устанавливаем по границе профиля
        # если есть участки 'Левая пойма', 'Правая пойма' задаем граиницы линии по участкам
        if x1 == None:
            x1 = min(self.morfostvor.x)
            for sector in self.morfostvor.sectors:
                if sector.name == 'Левая пойма':
                    x1 = sector.coord[0][-1]
        if x2 == None:
            x2 = max(self.morfostvor.x)
            for sector in self.morfostvor.sectors:
                if sector.name == 'Правая пойма':
                    x2 = sector.coord[0][0]

        # Подпись текста
        erosioin_limit_text = self.ax.text(x2 - 1, h + 0.01, text.format(
            h=h), color=config.COLOR['erosion_limit_text'], fontsize=config.FONT_SIZE['erosion_limit'], weight='bold')
        # Обводка текста
        erosioin_limit_text.set_path_effects([path_effects.Stroke(
            linewidth=3, foreground='white', alpha=0.95), path_effects.Normal()])

        # Отрисовка линии предельного размыва
        self.ax.plot([x1, x2], [h, h], color=config.COLOR['erosion_limit_line'],
                        linestyle='--', linewidth=config.LINE_WIDTH['erosion_limit_line'])
        # Добавляем в список границ отметку
        self._y_limits.append(h)
        self._update_limit()

    def draw_top_limit(self, h, x1=None, x2=None, text='{}\nH = {:.2f}'):
        y_step = self.ax.get_yticks()[1] - self.ax.get_yticks()[0]
        # Если координаты начала и конца линии не заданы, устанавливаем по границе профиля
        # если есть участки 'Левая пойма', 'Правая пойма' задаем граиницы линии по участкам
        if x1 == None:
            x1 = min(self.morfostvor.x)
            for sector in self.morfostvor.sectors:
                if sector.name == 'Левая пойма':
                    x1 = sector.coord[0][-1]
        if x2 == None:
            x2 = max(self.morfostvor.x)
            for sector in self.morfostvor.sectors:
                if sector.name == 'Правая пойма':
                    x2 = sector.coord[0][0]

        cent_x = x2 - ((x2 - x1) / 2)

        top_limit_text = self.ax.text(
            cent_x,
            h + (y_step * 0.2),
            '{}\nH = {:.2f}'.format(self.morfostvor.top_limit_description, h),
            color=config.COLOR['top_limit_text'],
            fontsize=config.FONT_SIZE['top_limit'],
            weight='bold',
            horizontalalignment='center',
            verticalalignment='center')
        # top_limit_text.set_path_effects([path_effects.Stroke(
        #     linewidth=3, foreground='white', alpha=0.95), path_effects.Normal()])
        # top_limit_text.set_path_effects([path_effects.Stroke(
            # linewidth=3, foreground='white', alpha=0.95), path_effects.Normal()])
        self.ax.plot([x1, x2], [h, h], color=config.COLOR['top_limit_line'],
                        linestyle='-.', linewidth=config.LINE_WIDTH['top_limit_line'])
        
        self._y_limits.append(h)
        self._update_limit()

    def draw_waterline(self, h, color=config.COLOR['water_line'], linestyle='--', linewidth=config.LINE_WIDTH['water_line']):
        """
        Функция отрисовки уреза воды по границам водного объекта.

        :param water: Исходный водный объект, содержащий координаты границ воды.
        :return: Отрисовывает урез на графике профиля (ax_profile).
        """

        def draw_line(self):
            for boundry in water.boundry():
                # Вводим служебые координаты
                x1, x2 = boundry[0][0], boundry[0][1]  # Начало и конец x
                y1, y2 = boundry[1][0], boundry[1][1]  # отметки уреза

                # Рисуем урез воды
                self.ax.plot([x1, x2], [y1, y2], color=color,
                                linestyle=linestyle, linewidth=linewidth)

                if config.PROFILE_WATER_FILL:
                    self.ax.fill(water.water_section_x, water.water_section_y,
                                    facecolor=config.COLOR['water_fill'], alpha=0.2)

        if config.PERELIV:
            water = WaterSection(self.morfostvor.x, self.morfostvor.y, h)
            draw_line(self)

        else:
            # Рисуем урезы на каждом участке
            for sector in self.morfostvor.sectors:
                x = sector.coord[0]
                y = sector.coord[1]

                if h >= min(y):
                    water = WaterSection(x, y, h)
                    draw_line(self)
        
        self._update_limit()
        self.set_style()

    def draw_levels_on_profile(self, levels):
        """
        Функция отрисовки полученных расчётных уровней воды на поперечном профиле.

        :param levels: DataFrame содержащий столбцы P, Q, H
        :return:
        """
        label = []

        for index, row in levels.iterrows():
            # Отрисовка уреза
            water_level = row['H']

            self.draw_waterline(water_level)

            if config.PROFILE_LEVELS_TITLE:
                # Подпись уровня воды на профиле
                water = WaterSection(self.morfostvor.x, self.morfostvor.y, water_level)
                try:
                    water = WaterSection(self.morfostvor.x, self.morfostvor.y, water_level)
                except:
                    print('Ошибка! При отрисовке расчётных уровней на профиле. \n')

                padding = 0.01
                x = water.water_section_x[0] + 2 * padding
                y = water_level + padding

                try:
                    # Если обеспеченность записана цифрами
                    waterline_text = self.ax.text(x, y, '▼$P_{{{:2g}\%}} = {:.2f}$'.format(
                        row['P'], row['H']), color=config.COLOR['water_level_text'], fontsize=config.FONT_SIZE['water_level'], weight='bold')
                    waterline_text.set_path_effects([path_effects.Stroke(
                        linewidth=3, foreground='white', alpha=0.55), path_effects.Normal()])
                except ValueError:
                    # Если обеспеченность записана строкой
                    waterline_text = self.ax.text(x, y, '{} = {:.2f}'.format(
                        row['P'], row['H']), color=config.COLOR['water_level_text'], fontsize=config.FONT_SIZE['water_level'], weight='bold')
                    waterline_text.set_path_effects([path_effects.Stroke(
                        linewidth=1.8, foreground='white', alpha=0.55), path_effects.Normal()])

            try:
                label.append('$P_{{{:2g}\%}} = {:.2f}$ м\n'.format(
                    row['P'], water_level))
            except ValueError:
                label.append('${} = {:.2f}$ м\n'.format(row['P'], water_level))

        if self.morfostvor.waterline and type(self.morfostvor.waterline) is not str:
            label.append('\nУВ = {:.2f} м\n'.format(self.morfostvor.waterline))
        
            if self.morfostvor.date:
                label.append('({})'.format(self.morfostvor.date))
        
        if config.PROFILE_WATER_LEVEL_NOTE:
            if self.morfostvor.waterline == '-' or self.morfostvor.waterline == '':
                label.append('\nПримечание: на\nмомент съёмки\nсток отсутствует')

        # Вывод таблицы уровней с разными обеспеченностями (справа)
        self.ax.annotate(''.join(label).rstrip(),
                                xy=(1.01, 0.8), xycoords='axes fraction',
                                size=config.FONT_SIZE['levels_table'],
                                color=config.COLOR['levels_table'],
                                bbox=dict(boxstyle="round", fc='white', ec="none"))

    def draw_wet_perimeter(self):
        """Функция отрисовки смоченного периметра на графике поперечного профиля
        """

        # Проверяем задан ли расчётный шаг в исходных данных
        if isinstance(self.morfostvor.dH, str) or self.morfostvor.dH == 0:
            self.morfostvor.dH = 1
            dH = self.morfostvor.dH
        else:
            dH = self.morfostvor.dH

        # Переводим сантиметры приращения в метры
        dH = dH / 100

        min_sector = self.morfostvor.get_min_sector()

        # Исходные сектора для расчёта (сектор, содержащий минимальную отметку)
        calc_sectors = [min_sector[0]]

        # Уровень воды, с минимальным отступом
        water_level = min(self.morfostvor.y) + dH


        # Цикл расчёта до максимальнjго уровня воды
        while water_level < self.morfostvor.gidraulic_result['УВ'].max():
            if config.PERELIV:
                for i in calc_sectors:
                    sector = self.morfostvor.sectors[i]
                    x = sector.coord[0]
                    y = sector.coord[1]

                    # Максимальная отметка слева
                    previus_min_ele = max(chunk_list(y, 2)[0])
                    # Максимальная отметка справа
                    next_min_ele = max(chunk_list(y, 2)[1])

                    # Проверка на перелив через границы участка
                    if (water_level >= previus_min_ele) and (i - 1 not in calc_sectors) and (i - 1 >= 0):
                        calc_sectors.append(i - 1)
                    if (water_level >= next_min_ele) and (i + 1 not in calc_sectors) and (i + 1 <= len(morfostvor.sectors) - 1):
                        calc_sectors.append(i + 1)

                    # Сектор воды и основные его параметры
                    # Расчетный участок является участком с минимальными отметками
                    # либо расчёт выполняется с одновременным заполнением
                    # начинаем заполнять с точки с минимальной отметкой
                    if sector.id == min_sector[1].id:
                        water = WaterSection(x, y, water_level)

                    # Расчетный участок находится слева от начального
                    # начинаем заполнять с крайней правой точки
                    elif sector.id < min_sector[1].id:
                        water = WaterSection(x, y, water_level, start_point=[
                                             len(y) - 1, y[-1]])

                    # Расчетный участок находится справа от начального
                    # начинаем заполнять с крайней левой точки
                    elif sector.id > min_sector[1].id:
                        water = WaterSection(
                            x, y, water_level, start_point=[0, y[0]])

                    # Отрисовка смоченного периметра на профиле на профиле
                    self.ax.plot(water.water_section_x, water.water_section_y,
                                            ':', marker='o', linewidth=1, color='black', markersize=3)
                    self.ax.plot([water.water_section_x[0], water.water_section_x[-1]], [
                                            water.water_section_y[0], water.water_section_y[-1]], ':', linewidth=1, color='black',)
            else:
                # Отрисовка с заполнением по участкам
                for sector in self.morfostvor.sectors:
                    x = sector.coord[0]
                    y = sector.coord[1]

                    if min(y) < water_level:
                        # Сектор воды и основные его параметры
                        water = WaterSection(x, y, water_level)

                        # Отрисовка смоченного периметра на профиле
                        self.ax.plot(water.water_section_x, water.water_section_y,
                                                ':', marker='o', linewidth=1, color='black', markersize=3)
                        self.ax.plot([water.water_section_x[0], water.water_section_x[-1]], [
                                                water.water_section_y[0], water.water_section_y[-1]], ':', linewidth=1, color='black',)

            water_level += dH

    def _update_limit(self):
        # Шаг засечек по вертикали
        y_step = self.ax.get_yticks()[1] - self.ax.get_yticks()[0]

        # Минимальное и максимальное значения из списка границ
        min_y = min(self._y_limits)
        max_y = max(self._y_limits)

        # Нижняя граница
        self.bottom_limit = np.ceil(min_y) - y_step
        while (self.morfostvor.ele_min - self.bottom_limit) < (y_step/3):
            self.bottom_limit -= y_step
        
        # Верхняя граница
        if (y_step > 0.5):
            self.top_limit = round(np.floor(max_y) + y_step, 3)
        else:
            self.top_limit = round((max_y // y_step * y_step) + y_step * 2, 3)         
        
        # Устанавливаем границы отображения   
        self._y_lim = (self.bottom_limit, self.top_limit)     
        self.ax.set_ylim(self._y_lim)
        self.draw_profile_point_lines()



def xls_calculate_hydraulic(in_filename, out_filename, page=None):
    """
    Выполнение гидравлических расчетов и создание отчета по результатам расчетов.
    Исходные данные представлены в in_filename (xls файл). По умолчанию расчеты производятся
    для всех листов xls файла. Если задан параметр page, расчет производится толька для указанной страницы.
    По результат создается out_filename (результирующий отчет в формате docx).

        :param in_filename: Входные данные по створам (.xls или .xlsx файл)
        :param out_filename: Результаты расчетов  (.docx файл)
        :param page=None: Номер страницы в xls файле, по умолчанию None (расчеты производятся для всего документа)
    """
    # Удаляем предыдущий отчет, если включена перезапись файла
    if config.REWRITE_DOC_FILE:
        try:
            os.remove(out_filename)
        except:
            pass

    page_quantity = get_xls_sheet_quantity(in_filename)
    stvors = []

    # Расчет для всех листов xls файла
    if page == None:
        for i in range(page_quantity):
            stvors.append(Morfostvor())
            stvors[i].read_xls(in_filename, i)
            stvors[i].calculate()
            stvors[i].doc_export(out_filename)

        # Вставка сводных таблиц
        insert_summary_QV_tables(stvors, out_filename)

    # Расчет только одного листа xls файла
    elif type(page) == int:
        stvor = Morfostvor()
        stvor.read_xls(in_filename, page)
        stvor.calculate()
        stvor.doc_export(out_filename)
    else:
        print('Номер листа должен быть int.')
        sys.exit(0)