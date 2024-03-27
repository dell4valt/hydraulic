import sys
from pathlib import Path
import numpy as np
from docx import Document
from hydraulic.doc_lib import (insert_page_break, set_table_style, set_table_columns_width)


def question_continue_app():
    while True:
        answer = input('Продолжить расчет? (да/нет)')
        if answer.lower() in ['да', 'д', 'yes', 'y', 'ага']:
            print('Хорошо. Продолжаем расчет.\n')
            break
        elif answer.lower() in ['no', 'нет', 'n', 'н']:
            print('Программа будет завершена.\n')
            sys.exit()
        else:
            continue


def poly_area(x, y):
    """
    Функция определения площади кривой фигуры.

        :param x: Список координат x
        :param y: Список координат y
    """
    return 0.5 * np.abs(np.dot(x, np.roll(y, 1)) - np.dot(y, np.roll(x, 1)))


def chunk_list(seq, num):
    """
    Функция разбивает заданный список на равные количество списков.

        :param seq: Исходный список
        :param num: Количество разбиваемых списков
    """
    avg = len(seq) / float(num)
    out = []
    last = 0.0

    while last < len(seq):
        out.append(seq[int(last): int(last + avg)])
        last += avg

    return out


def insert_summary_QV_tables(stvors, out_filename):
    print("Формируем и вставляем сводные таблицы уровней, "
          "скоростей и таблиц параметров при РУВВ... ", end="")
    # Подготовка данных для записи результирующей таблицы
    levels_table = []
    speed_table = []
    i = 1
    doc = Document(out_filename)

    insert_page_break(doc)

    doc.add_paragraph("Сводные таблицы", style="З-приложение-подзаголовок")

    param_levels = (
        ["№", "Описание", "Мин. отм", "УВ"],
        [0.85, 7, 1.25, 1.25],
        ["", "", ":.2f", ":.2f"],
    )

    param_speed = (
        [
            "№",
            "Описание",
        ],
        [
            0.85,
            8.5,
        ],
        [
            "",
            "",
        ],
    )

    title_check = False
    title2_check = False

    for stvor in stvors:
        levels_table.append([i, stvor.title, stvor.ele_min, stvor.waterline])
        speed_table.append(
            [
                i,
                stvor.title,
            ]
        )

        for obsp in stvor.levels_result.values.tolist():
            levels_table[i - 1].append(round(obsp[2], 2))

            if title_check is False:
                try:
                    param_levels[0].append("P{obsp[0]:g}%")
                except ValueError:
                    param_levels[0].append("{obsp[0]}")
                param_levels[1].append(1.25)
                param_levels[2].append(":.2f")

        for obsp in stvor.levels_result.values.tolist():
            speed_table[i - 1].append(round(obsp[3], 3))

            if title2_check is False:
                try:
                    param_speed[0].append(f"V{obsp[0]:g}%")
                except ValueError:
                    param_speed[0].append(f"{obsp[0]}")

                param_speed[1].append(1.25)
                param_speed[2].append(":.2f")

        title_check = True
        title2_check = True
        i += 1

    # rows = len(levels_table)
    cols = len(levels_table[0])

    #############################################
    # Таблицы расчётных уровней и скоростей воды
    doc.add_paragraph(
        f"Таблица — Расчётные уровни {stvor.strings['type']}",
        style="Т-название",
    )
    lev_table = doc.add_table(2, cols, style="Table Grid")

    doc.add_paragraph(
        f"Таблица — Расчётные скорости {stvor.strings['type']}",
        style="Т-название",
    )
    spd_table = doc.add_table(2, cols, style="Table Grid")

    doc.add_paragraph(
        "Таблица — Сводная таблица параметров РУВВ по поперечным профилям",
        style="Т-название",
    )
    ruvv_table = doc.add_table(1, 12, style="Table Grid")
    ruvv_table.cell(0, 0).text = '№'
    ruvv_table.cell(0, 1).text = '№ про-филя'
    ruvv_table.cell(0, 2).text = 'Описание'
    ruvv_table.cell(0, 3).text = 'Обеспе-ченность РУВВ'
    ruvv_table.cell(0, 4).text = 'Участок'
    ruvv_table.cell(0, 5).text = 'Уклон i, ‰'
    ruvv_table.cell(0, 6).text = 'Коэффициент шероховатости n'
    ruvv_table.cell(0, 7).text = 'Q при РУВВ, м³/сек'
    ruvv_table.cell(0, 8).text = 'Hср при РУВВ, м БС'
    ruvv_table.cell(0, 9).text = 'Vср при РУВВ, м/сек'
    ruvv_table.cell(0, 10).text = 'B при РУВВ, м'
    ruvv_table.cell(0, 11).text = 'F при РУВВ, м²'

    lev_table.cell(0, 0).merge(lev_table.cell(1, 0)).text = param_levels[0][0]
    lev_table.cell(0, 1).merge(lev_table.cell(1, 1)).text = param_levels[0][1]
    lev_table.cell(0, 2).merge(lev_table.cell(1, 2)).text = param_levels[0][2]
    lev_table.cell(0, 3).merge(lev_table.cell(1, 3)).text = param_levels[0][3]
    lev_table.cell(0, 4).merge(
        lev_table.cell(0, len(param_levels[0]) - 1)
    ).text = "Уровни воды (м БС), обеспеченностью Р%"

    spd_table.cell(0, 0).merge(spd_table.cell(1, 0)).text = param_levels[0][0]
    spd_table.cell(0, 1).merge(spd_table.cell(1, 1)).text = param_levels[0][1]
    spd_table.cell(0, 2).merge(spd_table.cell(1, 2)).text = param_levels[0][2]
    spd_table.cell(0, 3).merge(spd_table.cell(1, 3)).text = param_levels[0][3]
    spd_table.cell(0, 4).merge(
        spd_table.cell(0, len(param_levels[0]) - 1)
    ).text = "Скорости воды (м/с), обеспеченностью Р%"

    stvor_num = 1
    # Подписываем вероятности
    for i in range(len(stvors[0].probability)):
        try:
            lev_table.cell(1, i + 4).text = f"{stvors[0].probability[i][0]:g}"
            spd_table.cell(1, i + 4).text = f"{stvors[0].probability[i][0]:g}"
        except ValueError:
            lev_table.cell(1, i + 4).text = f"{stvors[0].probability[i][0]}"
            spd_table.cell(1, i + 4).text = f"{stvors[0].probability[i][0]}"

    # Заполняем сводные таблицы данными
    ruvv_n = 1
    for stvor in stvors:
        levels = stvor.levels_result[["P", "H", "Q"]].values.tolist()
        speed = stvor.levels_result[["P", "H", "V"]].values.tolist()

        lev_cell = lev_table.add_row().cells
        lev_cell[0].text = str(stvor_num)
        lev_cell[1].text = str(stvor.title)
        lev_cell[2].text = f"{stvor.ele_min:.2f}"

        spd_cell = spd_table.add_row().cells
        spd_cell[0].text = str(stvor_num)
        spd_cell[1].text = str(stvor.title)
        spd_cell[2].text = f"{stvor.ele_min:.2f}"

        # Проверка наличия уреза воды и вставка его в таблицу
        if isinstance(stvor.waterline, float):
            lev_cell[3].text = f"{stvor.waterline:.2f}"
            spd_cell[3].text = f"{stvor.waterline:.2f}"
        else:
            lev_cell[3].text = "-"
            spd_cell[3].text = "-"

        for i in range(4, len(levels) + 4):
            try:
                lev_cell[i].text = f"{levels[i - 4][1]:.2f}"
                spd_cell[i].text = f"{speed[i - 4][2]:.2f}"
            except:
                print(
                    "\n\nОшибка соответствия обеспеченностей в профилях.\
                     Обеспеченности на всех профилях должны быть одинаковые."
                )
                print("Сводные таблицы не будут записаны в файл.")

        sector_num = 1
        ruvv_cell = ruvv_table.add_row().cells

        for i in range(stvor.sectors_result.index.max() + 1):
            ruvv_cell[0].text = f"{ruvv_n}"
            ruvv_cell[4].text = f"{stvor.sectors_result.loc[i]['name']}"
            ruvv_cell[5].text = f"{stvor.sectors_result.loc[i]['slope']:.2f}".replace("nan", "-")
            ruvv_cell[6].text = (
                f"{stvor.sectors_result.loc[i]['roughness']:.3f}".replace("nan", "-")
            )
            ruvv_cell[7].text = (
                f"{stvor.sectors_result.loc[i]['consumption']:.2f}".replace("nan", "-")
            )
            ruvv_cell[8].text = f"{stvor.sectors_result.loc[i]['depth']:.2f}".replace("nan", "-")
            ruvv_cell[9].text = f"{stvor.sectors_result.loc[i]['speed']:.2f}".replace("nan", "-")
            ruvv_cell[10].text = f"{stvor.sectors_result.loc[i]['width']:.2f}".replace("nan", "-")
            ruvv_cell[11].text = f"{stvor.sectors_result.loc[i]['area']:.2f}".replace("nan", "-")
            sector_num += 1
            ruvv_cell = ruvv_table.add_row().cells
            ruvv_n += 1

        # Удаляем пустую ячейку
        row = ruvv_table.rows[-1]
        row._element.getparent().remove(row._element)

        # Объединяем ячейки
        prob_text = text_sanitize(
            stvor.probability[stvor.design_water_level_index][0], num_suffix="%"
        )
        ruvv_table.cell(ruvv_n - 1, 1).merge(
            ruvv_table.cell(ruvv_n - stvor.sectors_result.shape[0], 1)
        ).text = f"{stvor_num}"
        ruvv_table.cell(ruvv_n - 1, 2).merge(
            ruvv_table.cell(ruvv_n - stvor.sectors_result.shape[0], 2)
        ).text = f"{stvor.title}"
        ruvv_table.cell(ruvv_n - 1, 3).merge(
            ruvv_table.cell(ruvv_n - stvor.sectors_result.shape[0], 3)
        ).text = f"{prob_text}"
        stvor_num += 1

        set_table_style(lev_table)
        set_table_columns_width(
            lev_table,
            (
                0.85,
                7,
                1.25,
                1.25,
                1.25,
                1.25,
                1.25,
                1.25,
                1.25,
            ),
        )

        set_table_style(spd_table)
        set_table_columns_width(
            spd_table,
            (
                0.85,
                7,
                1.25,
                1.25,
                1.25,
                1.25,
                1.25,
                1.25,
                1.25,
            ),
        )

        set_table_style(ruvv_table)
        set_table_columns_width(
            ruvv_table,
            (
                0.85,
                0.85,
                3,
                1,
                2,
                2,
                2,
                2,
                2,
                2,
                2,
                2,
            ),
        )

    print("успешно!")
    doc.save(out_filename)


def text_sanitize(text, suffix='', prefix='', num_suffix=''):
    """Возвращает входной параметр text. В случае если число целое,
    возвращает без десятичных нулей. Если не целое, с указанием десятых.
    Можно задать префикс и суффикс соответствующими параметрами.

    Args:
        text (str, int, float): Входящая строка или число
        suffix (str, optional): Окончание возвращаемой строки. Defaults to ''.
        prefix (str, optional): Начало возвращаемой строки. Defaults to ''.
        num_suffix (str, optional): Окончание возвращаемой строки
        только если на входе число. Defaults to ''.

    Returns:
        _type_: Возвращаемая строка
    """

    try:
        return f'{prefix}{text:g}{num_suffix}{suffix}'
    except ValueError:
        return f'{prefix}{str(text)}{suffix}'


def rmdir(dir_path):
    directory = Path(str(dir_path))

    for item in directory.iterdir():
        if item.is_dir():
            rmdir(item)
        else:
            item.unlink()
    directory.rmdir()
