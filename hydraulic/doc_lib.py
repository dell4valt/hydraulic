"""Библиотека предоставляет набор полезных функция для
формирования технического отчета в формате .docx
с помощью библиотеки python-docx. Таких как:

    * insert_df_to_table - вставка таблицы из Pandas DataFrame
    * insert_chart - вставка графика Matplotlib в документ
    * insert_page_break - вставка разрыва страницы
    * set_table_columns_width - установка ширины колонок в таблице
    * set_table_style - установка стиля текста в ячейках таблицы
    * set_last_paragraph_style - установка стиля текста последнего параграфа
    * get_xls_sheet_quantity - получить количество листов в файле Excel
"""

import os
import random
import sys
from pathlib import Path

import xlrd
from docx.enum.text import WD_BREAK
from docx.shared import Cm


def insert_df_to_table(
    doc,
    df,
    title=None,
    footer_text=None,
    col_names=None,
    col_widths=None,
    col_format=None,
    table_style="Table Grid",
    text_style="Т-таблица",
    first_row_table_style="Т-таблица-заголовок"
):
    """
    Функция вставляет Pandas DataFrame в документ Word как отформатированную
    таблицу с возможностью указания заголовков для колонок, стилей таблицы и
    текста, ширины столбцов.

    Args:
        doc (docx.document.Document): Ссылка на объект документа Python-docx, в который вы хотите
    вставить таблицу.
        df (Pandas.DataFrame): DataFrame содержащий данные, которые необходимо вставить
    в таблицу документа Word.
        title (str): Параметр title используется для указания заголовка (названия) таблицы,
    которая будет вставлена в документ.
        col_names(tuple, list): Параметр переопределяет заголовки колонок в таблице.
    По умолчанию заголовки соответствуют названию колонок в DataFrame.
        col_widths (tuple): При указании, кортеж переопределяет ширину колонок таблицы по порядку.
        col_format (tuple): Параметр используется для указания стиля форматирования
    значений для каждого столбца таблицы.
    Форматы соответствует f-строке (пример: ":g", ":g", ":.2f")
        table_style (str): Название устанавливаемого стиля таблицы.
    Стиль должен присутствовать в файле шаблона отчета. По умолчанию "Table Grid".
        text_style: Название устанавливаемого основное стиля текста в таблице.
    Стиль должен присутствовать в файле шаблона отчета. По умолчанию "Т-таблица".
        first_row_table_style: Название устанавливаемого стиля строки заголовков таблицы.
    Стиль должен присутствовать в файле шаблона отчета. По умолчанию "Т-таблица-заголовок".

    Returns:
        Функция возвращает экземпляр docx.table.Table с произведенными изменениям.
    """
    # Вставляем заголовок таблицы
    if title:
        doc.add_paragraph(f"{title}", style="Т-название")

    # Количество строк и столбцов в таблице
    rows = df.shape[0]
    columns = df.shape[1]

    # Добавляем таблицу в документ
    table = doc.add_table(rows + 1, columns, style=table_style)

    # Получаем доступ к ячейкам экземпляра таблицы для
    # увеличения производительности, все последующие
    # операции производим с ячейками а не с экземпляром таблицы
    cells = table._cells

    # Устанавливаем 1ю строку заголовков
    for column_idx, column_name in enumerate(df.columns):
        if col_names:
            cells[column_idx].text = str(col_names[column_idx])
        else:
            cells[column_idx].text = str(column_name)

    # Записываем данные df в таблицу
    for row_idx in range(rows):
        for column_idx in range(columns):
            cell_value = df.iat[row_idx, column_idx]

            # Если задан список стилей для форматирования текста
            # устанавливаем формат для значения каждой ячейки
            # иначе просто записываем строку значения в ячейку
            if col_format and isinstance(cell_value, (float, int)):
                s = f"{{{col_format[column_idx]}}}"
                cells[column_idx + row_idx * columns + columns].text = s.format(cell_value)
            else:
                cells[column_idx + row_idx * columns + columns].text = str(cell_value)

    if col_widths:
        set_table_columns_width(table, col_widths)

    set_table_style(table, text_style, first_row_table_style)

    # Записываем служебный параграф после таблиц
    if footer_text:
        doc.add_paragraph(footer_text, style='Т-примечание')
    else:
        doc.add_paragraph('', style='Т-примечание')
    return table


def insert_figure(doc, chart, title='', dpi=200, width=16.5):
    """Функция вставляет график Matplotlib.plt в документ, предварительно
     сохранив его во временный файл. Устанавливает стили, добавляет заголовок
     и затем удаляет временный файл.

    Args:
      doc (docx.Document): Документ в который вставляется график
      chart (Matplotlib.plt): график который необходимо вставить в документ
      title (str): заголовок графика
    """

    # Временная директория в которую сохраняется график
    temp_dir = "TEMP"

    # Название временного файла графика
    filename = f'temp_{random.randrange(1, 10**3):03}.png'

    # Путь сохранения графика в файл
    file_path = f'{temp_dir}/{filename}'
    # Создаем временную папку, если она отсутствует
    Path(file_path).parents[0].mkdir(parents=True, exist_ok=True)
    # Сохраняем график в файл
    chart.savefig(file_path, dpi=dpi)
    # Вставляем график в документ отчёта
    doc.add_picture(file_path, width=Cm(width))
    # Устанавливаем стиль графика
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.style = 'Р-рисунок'
    doc.add_paragraph(
        f'{title}',
        style='Р-название')

    # Пытаемся удалить файл графика
    try:
        os.remove(file_path)
    except FileNotFoundError:
        pass


def insert_page_break(doc):
    """Процедура вставляет в документ Word разрыв страницы
    в том месте где вызывается.

    Args:
        doc (docx.Document): Файл документа в который вставляется
    разрыв страницы.
    """
    paragraphs = doc.paragraphs
    run = paragraphs[-1].add_run()
    run.add_break(WD_BREAK.PAGE)


def set_table_columns_width(table, col_widths: tuple):
    """Процедура устанавливает ширину столбцов в таблице docx
    поочередно проходя по каждой ячейке таблицы.

    Args:
        table (docx.table.Table): Таблица в документе docx.
        col_widths (tuple): Кортеж, содержащий ширину столбцов
    в сантиметрах по порядку.
    """
    # Получаем доступ к ячейкам таблицы из соображений производительности
    # и считаем количество колонок и строк
    cells = table._cells
    columns = len(table.columns)
    rows = len(table.rows)

    if columns != len(col_widths):
        print(
            "\nВнимание количество заданных столбцов "
            "не совпадает с количеством столбцов в таблице."
        )
        print(f"В таблице — {columns}, задано — {len(col_widths)}")

    for row_idx in range(rows):
        for column_idx in range(columns):
            # Номер ячейки по порядку
            cell_n = column_idx + row_idx * columns

            # Устанавливаем ширину столбцов.
            # Если ячеек заданных ширин меньше чем столбцов
            # устанавливаем последнее успешное значение
            try:
                success_width = Cm(col_widths[column_idx])
                cells[cell_n].width = success_width

            except IndexError:
                cells[cell_n].width = success_width


def set_table_style(table, style="Т-таблица", first_row_style="Т-таблица-заголовок"):
    """Процедура проходил по всем ячейкам таблицы table и устанавливает
    заданный стиль style параграфов в таблице. При желании можно
    указать стиль для заголовков таблицы (1-ая строка).

    Args:
        table (docx.table.Table): Таблица в документе docx
        style (str): Название устанавливаемого стиля. Стиль должен присутствовать
        в файле шаблона отчета. По умолчанию "Т-таблица".
        first_row_style (str): Названия стиля для первой строки таблицы.
        (строка заголовков). По умолчанию None.
    """
    # Получаем доступ к ячейкам таблицы из соображений производительности
    # и считаем количество колонок и строк
    cells = table._cells
    columns = len(table.columns)
    rows = len(table.rows)

    for row_idx in range(rows):
        for column_idx in range(columns):
            # Номер ячейки по порядку
            cell_n = column_idx + row_idx * columns

            # Если задан стиль первой строки устанавливаем
            # иначе проходим по всем ячейкам и устанавливаем
            # основной стиль таблицы
            for paragraph in cells[cell_n].paragraphs:
                paragraph.style = style

            if row_idx == 0 and first_row_style:
                for paragraph in cells[cell_n].paragraphs:
                    try:
                        paragraph.style = first_row_style
                    except TypeError:
                        print("Ошибка установки стиля первой строки, "
                              f"заголовок: {paragraph.text}, стиль: {first_row_style}")


def set_last_paragraph_style(doc, style: str):
    """Процедура устанавливает стиль последнего по порядку
    параграфа в документе Word.

    Args:
        doc (docx.Document): Файл документа в котором устанавливается стиль
    последнего по порядку параграфа.
        style (str): Название стиля в документе
    (стиль обязательно должен присутствовать в файле).
    """
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.style = style


def get_xls_sheet_quantity(file_path):
    """Функция возвращает количество листов в указанном xls файле.

    Args:
        file_path (str): Путь к xls файлу
    """
    try:
        # Открываем xls файл
        data_file = xlrd.open_workbook(file_path)
    except FileNotFoundError:
        print(
            "Ошибка определения количества листов в Excel файле! "
            f"Файл {file_path} не найден. Программа будет завершена."
        )
        sys.exit(33)

    return data_file.nsheets
