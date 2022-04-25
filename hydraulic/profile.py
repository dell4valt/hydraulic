import os
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

import matplotlib
import matplotlib.patheffects as path_effects
from matplotlib.patches import Rectangle
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import scipy.interpolate as interpolate
from setuptools import setup
import xlrd
from docx import Document
from docx.shared import Cm
from docxtpl import DocxTemplate
from matplotlib import gridspec
from pathvalidate import sanitize_filename

import hydraulic.config as config
from hydraulic.lib import (
    WD_BREAK,
    chunk_list,
    get_xls_sheet_quantity,
    insert_summary_QV_tables,
    insertPageBreak,
    poly_area,
    rmdir,
    setLastParagraphStyle,
    write_table,
)


@dataclass
class ProfileSector(object):
    """Класс участка профиля (пойма, русло и т.д.)

    :param id: Номер участка
    :param name: Описание (название) участка
    :param start_point: Номер первой точки участка
    :param end_point: Номер последней точки участка
    :param roughness: Коэффициент шероховатости n
    :param slope: Уклон данного участка I, ‰
    :param coord: Список с двумя подсписками координат x и y участка
    """

    id: int
    name: str
    start_point: int
    end_point: int
    roughness: float
    slope: float
    coord: tuple

    def __post_init__(self):
        self.color = self.get_color()

    def get_color(self):
        name = self.name
        channel = re.findall("русло", name, flags=re.IGNORECASE)
        protoka = re.findall("протока", name, flags=re.IGNORECASE)
        # floodplain = re.findall('пойма', name, flags=re.IGNORECASE)

        if channel:
            color = [0, 0.5, 1]
        elif protoka:
            color = [0, np.random.uniform(0, 0.5), np.random.uniform(0.5, 1)]
        # elif floodplain:
        # color = [np.random.uniform(.3, 1), 0, 0]
        else:
            color = [
                np.random.uniform(0, 1),
                np.random.uniform(0, 1),
                np.random.uniform(0, 1),
            ]
        return color

    def get_length(self):
        return round(self.coord[0][-1] - self.coord[0][0], 3)


@dataclass
class SituationSector(object):
    id: int
    type: str
    start_point: int
    end_point: int

@dataclass
class SituationBorder(object):
    id: int
    type: str



@dataclass
class WaterSection(object):
    """Класс водного сечения

    :param x: Точки x всего профиля
    :param y: Точки y всего профиля
    :param water_level: Уровень воды
    :param water_section_x: Точки x водного сечения
    :param water_section_y: Точки y водного сечения
    :param width: Ширина водного сечения
    :param area: Площадь водного сечения
    :param average_depth: Средняя глубина
    :param max_depth: Максимальная глубина
    :param wet_perimeter: Смоченный периметр
    :pararm r_hydraulic: Гидравлический радиус
    :param start_point: Точка начала расчёта [point_index, y]

    """

    x: float
    y: float
    water_level: float
    water_section_x: list = field(default_factory=list)
    water_section_y: list = field(default_factory=list)
    width: float = 0.0
    area: float = 0.0
    average_depth: float = 0.0
    max_depth: float = 0.0
    wet_perimeter: float = 0.0
    r_hydraulic: float = 0.0
    start_point: list = field(default_factory=list)

    def __post_init__(self):
        # start_point=[self.y.index(min(self.y)), min(self.y)]
        boundary = self.boundary()
        if len(boundary) > 1:
            for water_boundary in boundary:
                try:
                    self._calculate_parameters(water_boundary)
                except IndexError:
                    print(
                        "Ошибка в определении границ урезов! Программа будет завершена."
                    )
                    sys.exit(2)

            # Вычисления если урезов несколько
            self.width = sum(self.width)
            self.area = sum(self.area)
            self.average_depth = np.average(self.average_depth)
            self.max_depth = max(self.max_depth)
            self.wet_perimeter = sum(self.wet_perimeter)
            self.r_hydraulic = sum(self.r_hydraulic)

        else:
            try:
                self._calculate_parameters(boundary[0])
            except IndexError:
                print("Ошибка в определении границ урезов! Программа будет завершена.")
                sys.exit(2)

    def boundary(self):
        x = self.x
        y = self.y
        water_level = self.water_level  # Отметка уреза воды
        water_boundary_x, water_boundary_y, water_boundary_points = [], [], []
        result = []
        start_point = self.start_point

        if not start_point:
            start_point = [y.index(min(y)), min(y)]

        # Проверка на ошибку расположения уреза под поверхностью дна
        if water_level < min(y):
            print(
                "Ошибка! Уровень воды ниже низшей точки дна. Программа будет завершена с ошибкой."
            )
            sys.exit(1)
        else:
            # Цикл влево от стартовой точки
            for i in range(start_point[0], -1, -1):
                # Если индекс минимальной отметки совпадает с левой правой участка
                if start_point[0] == 0 and y[start_point[0]] <= water_level:
                    water_boundary_x.append(x[0])
                    water_boundary_y.append(water_level)
                    water_boundary_points.append(0)
                    break

                # Условие пересечения уреза с дном
                if y[i - 1] >= water_level and y[i] <= water_level:
                    x1, x2 = x[i - 1], x[i]
                    y1, y2 = y[i - 1], y[i]

                    # Нахождение координаты x уреза между точками дна
                    f = interpolate.interp1d([y1, y2], [x1, x2])
                    # Находим координату x, зная y (точка пересечения уреза с дном)
                    water_boundary_x.append(float(f(water_level)))
                    water_boundary_y.append(water_level)
                    # Присоединяем номер точки дна с границей воды
                    water_boundary_points.append(i - 1)
                    break  # Прерываем поиск если нашли пересечение

                # Условие отсутствия пересечения с дном и дохождения до начала участка
                elif i - 1 == 0 and y[i - 1] <= water_level:
                    water_boundary_x.append(x[i - 1])
                    water_boundary_y.append(water_level)
                    water_boundary_points.append(i - 1)
                    break  # Прерываем поиск если нашли пересечение

            # Цикл вправо от стартовой точки
            for i in range(start_point[0], len(y) - 1):
                # Условие пересечения уреза с дном
                if y[i] <= water_level and y[i + 1] >= water_level:
                    x1, x2 = x[i], x[i + 1]
                    y1, y2 = y[i], y[i + 1]

                    # Нахождение координаты x уреза между точками дна
                    f = interpolate.interp1d([y1, y2], [x1, x2])
                    # Находим координату x, зная y (точка пересечения уреза с дном)
                    water_boundary_x.append(float(f(water_level)))
                    water_boundary_y.append(water_level)
                    # Присоединяем номер точки дна с границей воды
                    water_boundary_points.append(i)
                    break  # Прерываем поиск если нашли пересечение

                elif i + 1 == len(y) - 1 and y[len(y) - 1] <= water_level:
                    water_boundary_x.append(x[len(x) - 1])
                    water_boundary_y.append(water_level)
                    water_boundary_points.append(i + 1)
                    break  # Прерываем поиск если нашли пересечение

            # Если индекс минимальной отметки совпадает с правой границой участка
            if start_point[0] == len(y) - 1 and y[start_point[0]] <= water_level:
                water_boundary_x.append(x[len(y) - 1])
                water_boundary_y.append(water_level)
                water_boundary_points.append(len(y) - 1)

            result.append(
                [water_boundary_x, water_boundary_y, water_boundary_points, 0]
            )
        return result

    # Функция выполняющая основные вычисления по данному водному сечению
    def _calculate_parameters(self, water_boundary):
        sum_sqr = 0
        water_level = self.water_level
        x = self.x
        y = self.y
        depth = []

        # Обрабатываем урезы по две точки (со второй до третьей пропускам)
        # Вводим служебные координаты (первая и последняя точки)
        x1, x2 = water_boundary[0][0], water_boundary[0][1]
        y1, y2 = water_boundary[1][0], water_boundary[1][1]

        # Точки смоченного периметра (номера точек под урезом)
        water_section_x = x[water_boundary[2][0] + 1 : water_boundary[2][1] + 1]
        water_section_y = y[water_boundary[2][0] + 1 : water_boundary[2][1] + 1]

        water_section_x.insert(0, x1)
        water_section_x.insert(len(water_section_x), x2)

        water_section_y.insert(0, y1)
        water_section_y.insert(len(water_section_y), y2)

        # Если первая точка УВ выше первой точки дна, вставляем точку дна на второе место
        # TODO: Костыль для определения полигона водной поверхности для расчёта с переливом
        #  и одновременным заполнением, нужно продумать как исправить
        if config.OVERFLOW:  # исходные данные точек x и y по всему профилю
            if water_level > y[water_boundary[2][0]]:
                water_section_x.insert(1, x[0])
                water_section_y.insert(1, y[0])
        else:  # исходные данные точек x и y по участкам
            if water_level > y[0]:
                water_section_x.insert(1, x[0])
                water_section_y.insert(1, y[0])

        # Если последняя точка УВ выше последней точки дна, вставляем точку на предпоследнее место
        if water_boundary[3] > 1 and water_level > y[-1]:
            water_section_x.insert(len(water_section_x) - 1, x[-1])
            water_section_y.insert(len(water_section_y) - 1, y[-1])

        # Координаты x и y смоченного периметра
        self.water_section_x = water_section_x
        self.water_section_y = water_section_y

        # Определяем ширину водной поверхности
        self.width = x2 - x1

        # Площадь воды
        self.area = poly_area(water_section_x, water_section_y)

        # Глубины
        for i in range(len(water_section_y)):
            depth.append(water_level - water_section_y[i])

        # Средняя глубина
        if self.area > 0 and self.width > 0:
            self.average_depth = self.area / self.width
        else:
            self.average_depth = 0

        if self.average_depth == 0:  # Костыль
            self.average_depth = 0.00001

        # Максимальная глубина
        self.max_depth = max(depth)

        # Смоченный периметр
        for i in range(len(water_section_x) - 1):
            sum_sqr += (water_section_x[i + 1] - water_section_x[i]) ** 2
        self.w_perimeter = np.sqrt(sum_sqr)

        # Гидравлический радиус
        if self.area > 0 and self.w_perimeter > 0:
            self.r_hydraulic = self.area / self.w_perimeter
        else:
            self.r_hydraulic = 0

        if self.r_hydraulic == 0:  # Костыль
            self.r_hydraulic = 0.00001


@dataclass
class Calculation(object):
    """
    Класс гидравлических расчётов скорости, расхода воды и коэффициента Шези для водного объекта.

    :param n: Коэффициент шероховатости
    :param i: Уклон, промилле
    :param h: Средняя глубина водного сечения
    :param a: Площадь водного сечения

    """

    n: float  # Коэффициент шероховатости
    i: float  # Уклон
    h: float  # Средняя глубина
    a: float  # Площадь водного сечений
    v: float = 0  # Скорость
    q: float = 0  # Расход
    __g: float = 9.80665  # Ускорение свободного падения
    shezi: float = 0  # Коэффициент Шези
    type__: str = "Не определен"

    def __post_init__(self):
        # В зависимости от глубины считаем по разным формулам
        # до 3-х метров по Павловскому, свыше 3-х метров по
        # Павловскому-Железнякову
        if self.h >= 0 and self.h <= 3:
            self.__shezi_pavlovskij()
        else:
            self.__shezi_pavlovskij_zheleznjakov()

        # Тип расчёта, обычная вода или селевой поток
        if config.CALC_TYPE == 1:
            # Расчёт скорости воды
            self.v = self.shezi * np.sqrt(self.h * (self.i / 1000))
        elif config.CALC_TYPE == 2:
            # Расчёт скорости воды для наносоводных селей
            self.v = 4.5 * self.h ** 0.67 * (self.i / 1000) ** 0.17
        elif config.CALC_TYPE == 3:
            # Расчёт скорости воды для грязекаменных селей селей
            self.v = 3.75 * self.h ** 0.50 * (self.i / 1000) ** 0.17
        else:
            print(
                "Ошибка выбора формулы расчёта скорости потока. Программа будет завершена."
            )
            sys.exit(1)
        # Расчёт расхода воды
        self.q = self.a * self.v

    # Коэффициент Шези по формуле Н. Н. Павловского, степенной коэффициент по формуле Железнякова
    def __shezi_pavlovskij_zheleznjakov(self):
        # Показатель сетепени по формуле Г. В. Железнякова
        y = (
            1
            / np.log10(self.h)
            * np.log10(
                (1 / 2 - (self.n * np.sqrt(self.__g) / 0.26) * (1 - np.log10(self.h)))
                + self.n
                * np.sqrt(
                    1
                    / 4
                    * (1 / self.n - np.sqrt(self.__g) / 0.13 * (1 - np.log10(self.h)))
                    ** 2
                    + np.sqrt(self.__g)
                    / 0.13
                    * (1 / self.n + np.sqrt(self.__g) * np.log10(self.h))
                )
            )
        )

        self.shezi = (1 / self.n) * self.h ** y
        self.type__ = "Коэффициент Шези определён по формуле Павловского, \
                       показатель степени определён по формуле Железнякова"

    # Коэффициент шези по формуле Маннинга
    def __shezi_mannign(self):
        self.shezi = (1 / self.n) * self.h ** (1 / 6)
        self.type__ = "Коэффициент Шези определён по формуле Маннинга"

    # Коэффициент шези по формуле Павловского для глубин 0.1 < h < 3 (Гидрорасчёты считают по этой формуле)
    def __shezi_pavlovskij(self):
        y = (
            2.5 * np.sqrt(self.n)
            - 0.13
            - 0.75 * np.sqrt(self.h) * (np.sqrt(self.n) - 0.10)
        )
        self.shezi = (1 / self.n) * self.h ** y
        self.type__ = (
            "Коэффициент шези определён по формуле Павловского для глубин 0.1 < h < 3 м"
        )

    # Коэффициент шези по формуле Железнякова
    def __shezi_zheleznjakov(self):
        self.shezi = 1 / 2 * (
            (1 / self.n) - (np.sqrt(self.__g) / 0.13) * (1 - np.log10(self.h))
        ) + np.sqrt(
            (1 / 4)
            * (1 / self.n - (np.sqrt(self.__g) / 0.13) * (1 - np.log10(self.h))) ** 2
            + (np.sqrt(self.__g) / 0.13)
            * ((1 / self.n) + (np.sqrt(self.__g) * np.log10(self.h)))
        )
        self.type__ = "Коэффициент шези определён по формуле Железнякова"


@dataclass
class Morfostvor(object):

    """Класс описывающий морфствор."""

    # Основные параметры морфоствора
    title: str = ""
    x: list = field(default_factory=list)
    y: list = field(default_factory=list)
    situation: list = field(default_factory=list)
    situation_borders: list = field(default_factory=list)
    sectors: list = field(default_factory=list)
    ele_max: float = 0
    ele_min: float = 0
    date: str = ""
    dH: int = 5
    waterline: float = 0
    erosion_limit: float = 0
    top_limit: float = 0
    top_limit_description: str = ""

    probability: list = field(default_factory=list)
    coords: list = field(default_factory=list)
    strings: dict = field(default_factory=dict)

    levels_result: pd.DataFrame = pd.DataFrame
    hydraulic_result: pd.DataFrame = pd.DataFrame
    hydraulic_table: pd.DataFrame = pd.DataFrame

    def __post_init__(self):
        # Выбор варианта расчёта
        if config.CALC_TYPE == 1:
            self.strings["type"] = "воды"
        elif config.CALC_TYPE == 2:
            self.strings["type"] = "наносоводного селевого потока"
        elif config.CALC_TYPE == 3:
            self.strings["type"] = "грязекаменного селевого потока"
        else:
            print(
                "Неверно выбран тип расчёта в конфигурационном файле. Программа будет завершена."
            )
            sys.exit(0)

        self.qh_title = f"Кривая расхода {self.strings['type']} Q = f(H)"

    def read_xls(self, file_path, page=0):
        """Функция чтения из xls файла."""
        # TODO: сделать проверку типа данных для коэффициента шероховатости

        try:
            data_file = xlrd.open_workbook(file_path)  # Открываем xls файл
        except FileNotFoundError:
            print(f"Ошибка! Файл {file_path} не найден. Программа будет завершена.")
            sys.exit(33)

        try:
            # Открываем лист по заданому номеру
            sheet = data_file.sheet_by_index(page)
        except IndexError:
            print(
                "Неверно указан индекс листа .xls файла. Проверьте параметры запуска расчёта."
            )
            sys.exit(1)

        sheet_title = sheet.name
        print(
            f"Считываем исходные данные из .xlsx файла: {file_path}, страница {page} ({sheet_title})."
        )

        __raw_data = []  # Сырые строки xlsx файла
        i = 0

        # Позиционирование столбцов с данными в .xls файле
        __x_coord_col = 0
        __y_coord_col = 1
        __sector_name_col = 2
        __roughness_col = 3
        __slope_col = 4
        __situation_col = 5
        __description_col = 8

        def get_situation(self):
            """Функция считывания участков ситуации из исходных файлов."""

            print("    — Определяем участки ситуации ... ", end="")

            lines_num = 0

            # Считываем количество строк с не пустыми координатами
            for line in __raw_data:
                if type(line[__x_coord_col]) != str:
                    lines_num += 1

            situation = self.situation
            situation_borders = self.situation_borders
            x = self.x  # Координаты профиля X
            num = 1

            for line in range(lines_num):
                try:
                    s1 = __raw_data[line][__situation_col].split(",")[0]
                    s2 = __raw_data[line][__situation_col].split(",")[1]

                    situation_borders.append(
                        SituationBorder(line, s2)
                    )
                except IndexError:
                    s1 = __raw_data[line][__situation_col]

                if line == 0:
                    situation.append(
                        SituationSector(num, s1, line, line)
                    )

                elif s1 != situation[num - 1].type:
                    if situation[num - 1].id == 1:
                        situation[num - 1].end_point = line
                    else:
                        situation[num - 1].end_point = line

                    num += 1

                    situation.append(
                        SituationSector(
                            num,
                            s1,
                            situation[num - 2].end_point,
                            line
                        )
                    )
            situation[-1].end_point = len(x) - 1

            print(f"успешно.\n")

            return(situation)

        def get_sectors(self):
            """Функция считывания участоков и их параметров из исходных файлов."""

            print("    — Определяем морфометрические участки ... ", end="")
            # №, Описание участка, номер первой точки, номер последней точки,
            # коэффициент шероховатости, уклон ‰, координата x, координаты y
            lines_num = 0

            # Считываем количество строк с не пустыми координатами
            for line in __raw_data:
                if type(line[__x_coord_col]) != str:
                    lines_num += 1

            sectors = self.sectors  # Список участков
            x = self.x  # Координаты профиля X
            y = self.y  # Координаты профиля Y

            num = 1  # Номера участков

            ###
            # Перебираем все строки xls файла и ищем участки
            for line in range(lines_num):
                name = __raw_data[line][__sector_name_col]  # Описание профиля
                # Коэффициент шероховатости
                roughness = __raw_data[line][__roughness_col]
                slope = __raw_data[line][__slope_col]  # Уклон

                # По первой строке создаём первый сектор
                if line == 0:
                    coord = []
                    sectors.append(
                        ProfileSector(num, name, line, line, roughness, slope, coord)
                    )

                # Сравниваем имя предыдущего участка с текущим, и если не совпадают то создаем новый сектор:
                elif name != sectors[num - 1].name:

                    # TODO: Проверить это условие
                    if sectors[num - 1].id == 1:  # Если первый участок
                        # Записываем номер последний точки - 1
                        sectors[num - 1].end_point = line
                    else:  # Если все остальные участки
                        # Записываем номер последний точки в предыдущий участок для всех остальных участков
                        sectors[num - 1].end_point = line

                    num += 1  # Увеличиваем номер сектора на 1
                    sectors.append(
                        ProfileSector(
                            num,
                            name,
                            sectors[num - 2].end_point,
                            line,
                            roughness,
                            slope,
                            coord,
                        )
                    )

            # Номер последней точки в последнем секторе
            sectors[-1].end_point = len(x) - 1

            # Записываем координаты и длины участков
            for sector in sectors:
                sector.coord = (
                    x[sector.start_point : sector.end_point + 1],
                    y[sector.start_point : sector.end_point + 1],
                )  # Координаты из начальной и конечной точек
                # Длины полученные из разницы координат по x
                sector.length = sector.get_length()

            try:
                # Максимальная отметка участка слева
                self.max_l = max(chunk_list(sector.coord[1], 2)[0])
                # Максимальная отметка участка справа
                self.max_r = max(chunk_list(sector.coord[1], 2)[1])
            except:
                print("\n\nОшибка в определении участков. Список участков:\n")
                for sector in sectors:
                    print(sector)

                print("Завершаем программу.")
                sys.exit(3)

            print(f"успешно, найдено {len(sectors)} участка.")
            return sectors

        # Перебираем все строки
        # И получаем список сырых данных
        for rownum in range(1, sheet.nrows):
            row = sheet.row_values(rownum)
            __raw_data.append(row)  # Записываем данные

        # Устанавливаем основные параметры морфоствора
        print("    — Устанавливаем основные параметры морфоствора ... ", end="")
        self.title = __raw_data[2][__description_col]  # Заголовок профиля
        self.date = __raw_data[3][__description_col]  # Дата профиля

        self.waterline = __raw_data[4][__description_col]  # Отметка уреза воды
        # Проверяем задан ли урез текстом, если нет округляем до 2 знаков
        if type(self.waterline) is not str:
            self.waterline = round(self.waterline, 2)

        self.dH = __raw_data[5][__description_col]  # Расчётный шаг по глубине
        self.coords = __raw_data[6][__description_col]  # Координаты
        self.erosion_limit = __raw_data[7][__description_col]  # Предел размыва
        self.top_limit = __raw_data[8][__description_col]  # Верхняя граница
        self.top_limit_description = __raw_data[9][
            __description_col
        ]  # Описание верхней границы
        print("успешно!")

        # Считываем и записываем все точки x и y профиля
        print("    — Считываем координаты профиля ... ", end="")
        for i in range(len(__raw_data)):
            if type(__raw_data[i][__x_coord_col]) != str:
                self.x.append(__raw_data[i][__x_coord_col])
                self.y.append(__raw_data[i][__y_coord_col])
                #self.situation.append(__raw_data[i][__situation_col])
        print(f"успешно, найдено {len(self.x)} точки, длина профиля {self.x[-1]:.2f} м")

        self.ele_min = min(self.y)  # Минимальная отметка профиля
        self.ele_max = max(self.y)  # Максимальная отметка профиля

        # Заполнения таблицы обеспеченностей
        print("    — Считываем обеспеченности ... ", end="")
        for i in range(6, len(__raw_data[0])):
            self.probability.append([__raw_data[0][i], __raw_data[1][i]])

        # Удаляем пустые обеспеченности из списка обеспеченностей
        self.probability = [x for x in self.probability if x != ["", ""]]

        print(f"успешно, найдено {len(self.probability)} обеспеченностей.")

        # Обработка и получение данных по секторам из "сырых" данных
        self.sectors = get_sectors(self)
        self.situation = get_situation(self)

    def get_min_sector(self):
        """
        Функция нахождения участка с наименьшей отметкой дна.

        :return: [Номер по списку, [Участок]]
        """

        id = 0
        i = 0
        min_sector = self.sectors[0]

        for sector in self.sectors:
            if min(sector.coord[1]) < min(min_sector.coord[1]):
                min_sector = sector
                id = i
            i += 1
        return (id, min_sector)

    def get_q_max(self):
        """
        Функция нахождения максимальной обеспеченности и расхода воды по исходным данным.

        :return: [Обеспеченность, Расход]
        """
        q_max = float(self.probability[0][1])
        obsp = self.probability[0][0]
        for Q in self.probability:
            if q_max <= Q[1]:
                q_max = Q[1]
                obsp = Q[0]

        return (obsp, q_max)

    def doc_export(self, out_filename, r=False):
        print("\n\nФормируем doc файл: ")
        doc_file = out_filename
        template_file = Path("hydraulic/assets/report_template.docx")

        # Создаем временную папку, и папку для графики если они не существуют
        temp_dir = Path(config.TEMP_DIR_NAME)
        temp_dir.mkdir(parents=True, exist_ok=True)

        # Создаем папку для сохранения отдельных изображений
        if config.PROFILE_SAVE_PICTURES:
            picture_dir = Path(
                str(Path(out_filename).parents[0]) + "/" + config.GRAPHICS_DIR_NAME
            )
            picture_dir.mkdir(parents=True, exist_ok=True)

        if r:
            doc = DocxTemplate(template_file)
        else:
            if os.path.isfile(doc_file):
                doc = Document(doc_file)
                paragraphs = doc.paragraphs
                run = paragraphs[-1].add_run()
                run.add_break(WD_BREAK.PAGE)
            else:
                if config.REWRITE_DOC_FILE:
                    print(
                        "    — Включена перезапись файла, удаляем старый и создаём новый."
                    )
                else:
                    print("    — Файл не найден! Создаём новый.")
                doc = DocxTemplate(template_file)

        if config.HYDRAULIC_CURVE:
            self.fig_QH = GraphQH(self)
        if config.SPEED_CURVE:
            self.fig_QV = GraphQV(self)
        if config.AREA_CURVE:
            self.fig_QF = GraphQF(self)

        self.fig_profile = GraphProfile(self)

        # Отрисовка смоченного периметра
        if config.PROFILE_WET_PERIMITER:
            self.fig_profile.draw_wet_perimeter()

        # Отрисовка верхней границы сооружения
        if self.top_limit:
            self.fig_profile.draw_top_limit(
                self.top_limit, text=self.top_limit_description
            )

        # Отрисовка границы предельного размыва профиля
        if self.erosion_limit:
            self.fig_profile.draw_erosion_limit(self.erosion_limit)

        # Отрисовка уровней воды на графике профиля
        self.fig_profile.draw_levels_on_profile(self.levels_result)

        self.fig_profile._update_limit()
        if self.waterline and type(self.waterline) != str:
            self.fig_profile.draw_waterline(
                round(self.waterline, 2), color="blue", linestyle="-"
            )

        print("    — Сохраняем график профиля ... ", end="")
        self.fig_profile.fig.savefig(
            Path(f"{config.TEMP_DIR_NAME}/Profile.png", dpi=config.FIG_DPI)
        )
        print("успешно!")

        # Вставляем заголовок профиля
        doc.add_paragraph(self.title, style="З-приложение-подзаголовок")

        # Добавляем изображения профиля и гидравлической кривой
        print("    — Вставляем графику (профиль и кривую)... ", end="")
        doc.add_picture(f"{config.TEMP_DIR_NAME}/Profile.png", width=Cm(16.5))
        setLastParagraphStyle("Р-рисунок", doc)

        # Подпись рисунков
        if config.GRAPHICS_TITLES_TEXT:
            doc.add_paragraph(
                "Рисунок — " + self.fig_profile.morfostvor.title, style="Р-название"
            )

        print("успешно!")

        if config.HYDRAULIC_CURVE:
            print("    — Сохраняем график гидравлической кривой ... ", end="")
            self.fig_QH.fig.savefig(
                Path(f"{config.TEMP_DIR_NAME}/QH.png", dpi=config.FIG_DPI)
            )
            print("успешно!")

            doc.add_picture(f"{config.TEMP_DIR_NAME}/QH.png", width=Cm(16.5))
            setLastParagraphStyle("Р-рисунок", doc)

            if config.GRAPHICS_TITLES_TEXT:
                doc.add_paragraph(
                    "Рисунок — " + self.fig_QH._ax_title_text, style="Р-название"
                )

        # Вставляем разрыв страницы
        insertPageBreak(doc)

        if config.SPEED_CURVE:
            print("    — Сохраняем график кривой скоростей ... ", end="")
            self.fig_QV.fig.savefig(
                Path(f"{config.TEMP_DIR_NAME}/QV.png", dpi=config.FIG_DPI)
            )
            print("успешно!")

            doc.add_picture(f"{config.TEMP_DIR_NAME}/QV.png", width=Cm(16.5))
            setLastParagraphStyle("Р-рисунок", doc)
            print("успешно!")

            if config.GRAPHICS_TITLES_TEXT:
                doc.add_paragraph(
                    "Рисунок — " + self.fig_QV._ax_title_text, style="Р-название"
                )

        if config.AREA_CURVE:
            print("    — Сохраняем график кривой площадей ... ", end="")
            self.fig_QF.fig.savefig(
                Path(f"{config.TEMP_DIR_NAME}/QF.png", dpi=config.FIG_DPI)
            )
            print("успешно!")

            doc.add_picture(f"{config.TEMP_DIR_NAME}/QF.png", width=Cm(16.5))
            setLastParagraphStyle("Р-рисунок", doc)
            print("успешно!")

            if config.GRAPHICS_TITLES_TEXT:
                doc.add_paragraph(
                    "Рисунок — " + self.fig_QF._ax_title_text, style="Р-название"
                )

        # Проверяем имя файла
        profile_name = sanitize_filename(self.title)

        # Сохраняем картинки в отдельные файлы в папку graphics
        if config.PROFILE_SAVE_PICTURES:
            self.fig_profile.fig.savefig(
                Path(f"{picture_dir}/{profile_name}.png", dpi=config.FIG_DPI)
            )
        if config.CURVE_SAVE_PICTURES:
            self.fig_QH.fig.savefig(
                Path(f"{picture_dir}/{profile_name}_QH.png", dpi=config.FIG_DPI)
            )
            self.fig_QV.fig.savefig(
                Path(f"{picture_dir}/{profile_name}_QV.png", dpi=config.FIG_DPI)
            )
            self.fig_QF.fig.savefig(
                Path(f"{picture_dir}/{profile_name}_QF.png", dpi=config.FIG_DPI)
            )

        # Вывод таблицы расчётных уровней воды
        print("    — Записываем таблицу уровней воды ... ", end="")
        param = (
            ("Обеспеченность P, %", "Расход Q, м³/сек", "Уровень H, м"),
            (5.6, 5.6, 5.6),
            (":g", ":g", ":.2f"),
        )
        print("успешно!")

        # levels_result = self.levels_result
        levels_result = self.levels_result[["P", "Q", "H"]].round(3).values.tolist()
        write_table(
            doc,
            levels_result,
            param,
            f"Таблица - Расчётные уровни {self.strings['type']}",
        )

        # Вывод таблицы участков
        print("    — Записываем таблицу участков ... ", end="")
        param = (
            ("№", "Описание", "Уклон i, ‰", "Коэффициент шероховатости n"),
            (1.5, 5.1, 5.1, 5.1),
            (":d", "", ":g", ":.3f"),
        )
        sectors = []

        i = 1
        for sector in self.sectors:
            sectors.append([i, sector.name, sector.slope, sector.roughness])
            i += 1

        write_table(doc, sectors, param, "Таблица - Расчётные участки и их параметры")
        print("успешно!")

        print("    — Записываем таблицу кривой расхода воды ... ", end="")

        # Вывод таблицы гидравлической кривой
        param = (
            (
                "Отм. уровня H, м",
                "Площадь F, м²",
                "Ширина B, м",
                "Средняя глубина Hср, м",
                "Макс. глубина Hмакс, м",
                "Средняя скорость Vср, м/сек",
                "Расход Q, м³/сек",
            ),
            (5, 5, 5, 5, 5, 5, 5),
            (":.2f", ":.3f", ":.3f", ":.3f", ":.3f", ":.3f", ":.3f"),
        )

        table = self.hydraulic_table.reset_index(0).loc["Сумма"].reset_index(drop=True)
        table_round = table.round(3)  # Округляем
        # Убираем столбец с коэффициентами Шези
        table_round = table_round.drop(columns=["Shezi"])

        if config.DOC_TABLE_SHORT:
            # Количество строк в таблице
            table_quant = table_round["УВ"].count()

            # Уменьшаем количество выводимых строк в таблицу
            # чтобы поместилось на один лист
            if table_quant <= 25:
                divider = 1
            elif table_quant > 25 and table_quant <= 50:
                divider = 2
            elif table_quant > 50 and table_quant <= 75:
                divider = 3
            elif table_quant > 75 and table_quant <= 100:
                divider = 4
            elif table_quant > 100 and table_quant <= 125:
                divider = 5
            else:
                divider = 10
        else:
            divider = 1

        # Записываем только чётные элементы таблицы
        table_round = table_round[table_round.index % divider == 0]
        sum_hydraulic = table_round.values.tolist()  # Переводим в список
        write_table(
            doc,
            sum_hydraulic,
            param,
            f"Таблица - Параметры расчёта кривой расхода {self.strings['type']}",
        )
        doc.add_paragraph(
            f"Расчётный шаг: {self.dH:g} см. В таблице приведён каждый {divider}-й результат расчёта.",
            style="Т-примечание",
        )
        print("успешно!")

        # Удаляем объект профиля
        self.fig_profile.clean()

        try:
            doc.save(doc_file)
        except PermissionError:
            print(
                "\nОшибка! Не удалось сохранить файл. Проверьте возможность записи файла по указанному пути."
            )
            print("Возможно записываемый файл уже существует и открыт.")
            sys.exit(1)

        # Удаляем временную папку со всем содержимым
        print("    — Удаляем временную папку ... ", end="")
        rmdir(Path(f"{config.TEMP_DIR_NAME}"))
        print("успешно!")

        print(
            f"\nФайл {doc_file} сохранён успешно.\n-------------------------------------\n"
        )

    def calculate(self):
        # Значение расхода до которого необходимо считать (максимальной введенная обеспеченности + 20%)
        consumption_check = self.get_q_max()[1] + (self.get_q_max()[1] * 0.20)

        # Проверяем задан ли расчётный шаг в исходных данных
        if isinstance(self.dH, str) or self.dH == 0:
            self.dH = 1
            dH = self.dH
        else:
            dH = self.dH

        # Переводим сантиметры приращения в метры
        dH = dH / 100

        min_sector = self.get_min_sector()

        # Исходные сектора для расчёта (сектор, содержащий минимальную отметку)
        calc_sectors = [min_sector[0]]

        # Уровень воды, с минимальным отступом
        water_level = min(self.y) + dH

        # Обнулённые переменные
        consumption_summ = 0
        area_summ = 0
        n = 0

        col = ["Участок", "УВ", "F", "B", "Hср", "Hмакс", "V", "Q", "Shezi"]
        df = pd.DataFrame(columns=col, dtype=float)
        # Первый расчётный элемент суммирующей кривой со всеми нулями§
        df = df.append(
            dict(zip(col, ["Сумма", self.ele_min, 0, 0, 0, 0, 0, 0, 0])),
            ignore_index=True,
        )

        # Цикл расчёта до максимальной обеспеченности + 20% из исходных данных
        while consumption_summ < consumption_check:
            print(f"Выполняем расчёты для уровня {water_level:.2f}", end="\r")

            consumption_summ = 0
            wc_list = list()
            area_list = list()

            if config.OVERFLOW:
                for i in calc_sectors:
                    sector = self.sectors[i]
                    x = sector.coord[0]
                    y = sector.coord[1]

                    # Максимальная отметка слева
                    previous_min_ele = max(chunk_list(y, 2)[0])
                    # Максимальная отметка справа
                    next_min_ele = max(chunk_list(y, 2)[1])

                    # Проверка на перелив через границы участка
                    if (
                        (water_level >= previous_min_ele)
                        and (i - 1 not in calc_sectors)
                        and (i - 1 >= 0)
                    ):
                        calc_sectors.append(i - 1)
                    if (
                        (water_level >= next_min_ele)
                        and (i + 1 not in calc_sectors)
                        and (i + 1 <= len(self.sectors) - 1)
                    ):
                        calc_sectors.append(i + 1)

                    # Сектор воды и основные его параметры
                    # Расчетный участок является участком с минимальными отметками
                    # либо расчёт выполняется с одновременным заполнением
                    # начинаем заполнять с точки с минимальной отметкой
                    if sector.id == min_sector[1].id:
                        water = WaterSection(x, y, water_level)

                    # Расчетный участок находится слева от начального
                    # начинаем заполнять с крайней правой точки
                    elif sector.id < min_sector[1].id:
                        water = WaterSection(
                            x, y, water_level, start_point=[len(y) - 1, y[-1]]
                        )

                    # Расчетный участок находится справа от начального
                    # начинаем заполнять с крайней левой точки
                    elif sector.id > min_sector[1].id:
                        water = WaterSection(x, y, water_level, start_point=[0, y[0]])

                    # Расчёт параметров для воды
                    calc = Calculation(
                        h=water.average_depth,
                        n=sector.roughness,
                        i=sector.slope,
                        a=water.area,
                    )

                    wc_list.append(calc.q)

                    r = dict(
                        zip(
                            col,
                            [
                                sector.name,
                                round(water_level, 2),
                                water.area,
                                water.width,
                                water.average_depth,
                                water.max_depth,
                                calc.v,
                                calc.q,
                                calc.shezi,
                            ],
                        )
                    )

                    # Добавляем в список с результирующими значениями значения по секторам
                    # для последующего суммирования/вычисления средних значений
                    df = df.append(r, ignore_index=True)

            else:
                # Расчёт с заполнением по участкам
                for sector in self.sectors:
                    x = sector.coord[0]
                    y = sector.coord[1]

                    if min(y) < water_level:
                        # Сектор воды и основные его параметры
                        water = WaterSection(x, y, water_level)

                        # Расчёт параметров для воды
                        calc = Calculation(
                            h=water.average_depth,
                            n=sector.roughness,
                            i=sector.slope,
                            a=water.area,
                        )

                        wc_list.append(calc.q)

                        # Добавляем в список с значения по секторам
                        r = dict(
                            zip(
                                col,
                                [
                                    sector.name,
                                    round(water_level, 2),
                                    water.area,
                                    water.width,
                                    water.average_depth,
                                    water.max_depth,
                                    calc.v,
                                    calc.q,
                                    calc.shezi,
                                ],
                            )
                        )

                        # Добавляем в список с результирующими значениями значения по секторам
                        # для последующего суммирования/вычисления средних значений
                        df = df.append(r, ignore_index=True)

            consumption_summ += sum(wc_list)
            area_summ += sum(area_list)

            # Пустые значения для суммирующей кривой
            r_sum = dict(
                zip(col, ["Сумма", round(water_level, 2), 0, 0, 0, 0, 0, 0, 0])
            )
            df = df.append(r_sum, ignore_index=True)

            water_level += dH
            n += 1

        # TODO: remake to use one dataframe
        df = df.set_index(["УВ", "Участок"])
        water_levels = df.index.levels[0]

        # Заполняем суммирующие данные
        df.loc[(water_levels, "Сумма"), "F"] = df.groupby(level=0)["F"].transform("sum")
        df.loc[(water_levels, "Сумма"), "B"] = df.groupby(level=0)["B"].transform("sum")
        df.loc[(water_levels, "Сумма"), "Hср"] = df.groupby(level=0)["F"].transform(
            "sum"
        ) / df.groupby(level=0)["B"].transform("sum")
        df.loc[(water_levels, "Сумма"), "Hмакс"] = df.groupby(level=0)[
            "Hмакс"
        ].transform("max")
        df.loc[(water_levels, "Сумма"), "Q"] = df.groupby(level=0)["Q"].transform("sum")
        df.loc[(water_levels, "Сумма"), "V"] = df.groupby(level=0)["Q"].transform(
            "sum"
        ) / df.groupby(level=0)["F"].transform("sum")
        df.loc[(water_levels, "Сумма"), "Shezi"] = df.groupby(level=0)[
            "Shezi"
        ].transform("sum") / (df.groupby(level=0)["Shezi"].transform("count") - 1)
        df = df.fillna(0)

        # Интерполируем значения гидравлической кривой
        # для необходимых обеспеченностей и обновляем таблицу
        p_table = df.loc[(water_levels, "Сумма"), :].droplevel(1)
        self.levels_result = self.get_p_table(p_table)
        self.hydraulic_table = df

        return df

    def get_p_table(self, df: pd.DataFrame):
        result = pd.DataFrame(columns=["P", "Q", "H", "F"])

        for prob in self.probability:
            fQ = interpolate.interp1d(df["Q"], df.index)
            fV = interpolate.interp1d(df["Q"], df["V"])
            fF = interpolate.interp1d(df["Q"], df["F"])
            h = float(fQ(prob[1]))
            v = float(fV(prob[1]))
            f = float(fF(prob[1]))

            result = result.append(
                {"P": prob[0], "H": h, "Q": prob[1], "V": v, "F": f}, ignore_index=True
            )

        return result


@dataclass
class Graph(object):
    _fig_size = (16.5, 11)
    _y_limits = []
    _fig_num = 0

    _x_label_text = ""
    _y_label_text = ""
    _ax_title_text = ""

    morfostvor: Morfostvor = Morfostvor
    fig: plt.figure = plt.figure(_fig_num, figsize=_fig_size)
    ax: plt.subplot = fig.add_subplot(111)

    def __post_init__(self):
        self.clean()
        morfostvor = self.morfostvor

        # Вытягиваем цвета
        self.sector_colors = {}
        for sector in morfostvor.sectors:
            self.sector_colors[sector.name] = sector.color

        # Выполняем отрисовку содержимого
        self.draw()
        self.set_style()

    def draw(self):
        pass

    def set_style(self):
        fig = self.fig
        ax = self.ax

        fig.subplots_adjust(bottom=0.08, left=0.08, right=0.9)

        # Устанавливаем заголовки графиков
        if config.GRAPHICS_TITLES:
            ax.set_title(
                self._ax_title_text,
                color=config.COLOR["title_text"],
                fontsize=config.FONT_SIZE["title"],
                y=1.1,
            )

        # Настраиваем границы и толщину линий границ
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_linewidth(config.LINE_WIDTH["ax_border"])
        ax.spines["bottom"].set_linewidth(config.LINE_WIDTH["ax_border"])

        # Устанавливаем отступы в графиках
        ax.margins(0.025)

        # Устанавливаем параметры засечек на основных осях
        ax.tick_params(
            which="major",
            direction="out",
            width=2,
            length=5,
            pad=10,
            labelcolor=config.COLOR["ax_value_text"],
            labelsize=config.FONT_SIZE["ax_major"],
        )

        ax.tick_params(
            which="minor",
            direction="out",
            width=1.5,
            length=3.5,
            pad=10,
            labelcolor=config.COLOR["ax_value_text"],
            labelsize=config.FONT_SIZE["ax_minor"],
        )

        # Устанавливаем параметры подписей осей
        ax.set_xlabel(
            self._x_label_text,
            color=config.COLOR["ax_label_text"],
            fontsize=config.FONT_SIZE["ax_label"],
            fontstyle="italic",
        )
        ax.xaxis.set_label_coords(1.05, -0.025)
        ax.set_ylabel(
            self._y_label_text,
            color=config.COLOR["ax_label_text"],
            fontsize=config.FONT_SIZE["ax_label"],
            fontstyle="italic",
        )
        ax.yaxis.set_label_coords(-0.025, 1.08)

        # Устанавливает параметры вывода значений осей
        ax.yaxis.set_major_formatter(matplotlib.ticker.FormatStrFormatter("%.10g"))

        # Настройка параметров отображение сетки
        ax.grid(
            which="major",
            color=config.COLOR["ax_grid"],
            linestyle=":",
            linewidth=1,
            alpha=0.9,
        )
        ax.grid(
            which="minor",
            color=config.COLOR["ax_grid_sub"],
            linestyle=":",
            linewidth=1,
            alpha=0.9,
        )

    def clean(self):
        """Очистка осей графика и обнуление связанных переменных"""
        # Очищаем все оси
        for ax in vars(self):
            if ax.startswith("ax"):
                command = "self." + ax + ".cla()"
                exec(command)

        # Обнуляем границы y
        self._y_limits = []
        self._y_limits = []


@dataclass
class GraphCurve(Graph):
    def draw_water_levels(
        self, morfostvor: Morfostvor, ax: plt.subplot, x="Q", y="H", y_min=0
    ):
        """Функция выводит на график ax отметку и линии пересечения
           x и y.

        Args:
            morfostvor (Morfostvor): Объект из которого необходимо брать данные.
            ax (plt.subplot): График для нанесения отметок.
            x (str, optional): Ось x. Defaults to 'Q'.
            y (str, optional): Ось y. Defaults to 'H'.
        """
        try:
            if config.HYDRAULIC_CURVE_LEVELS:
                for index, row in morfostvor.levels_result.iterrows():
                    x1, x2 = 0, row[x]
                    y1, y2 = row[y], row[y]

                    # Вывод значений округленных, проверка на содержание значений
                    try:
                        water_level_text = ax.text(
                            0.002,
                            row[y],
                            f"▼$P_{{{row['P']:.2g}\\%}} = {row[y]:.2f}$",
                            color=config.COLOR["water_level_text"],
                            fontsize=config.FONT_SIZE["water_level"],
                            weight="bold",
                        )

                        water_level_text.set_path_effects(
                            [
                                path_effects.Stroke(
                                    linewidth=3, foreground="white", alpha=0.55
                                ),
                                path_effects.Normal(),
                            ]
                        )

                    except ValueError:
                        water_level_text = ax.text(
                            0.002,
                            row[y],
                            f"▼${row['P']} = {row[y]:.2f}$",
                            color=config.COLOR["water_level_text"],
                            fontsize=config.FONT_SIZE["water_level"],
                            weight="bold",
                        )

                        water_level_text.set_path_effects(
                            [
                                path_effects.Stroke(
                                    linewidth=3, foreground="white", alpha=0.55
                                ),
                                path_effects.Normal(),
                            ]
                        )

                    ax.plot(
                        [x1, x2, x2, x2],
                        [y1, y2, y_min, y_min],
                        linestyle="-",
                        color="mediumturquoise",
                        marker="o",
                        linewidth=1,
                        markersize=1,
                    )
        except:
            print("Внимание! Вывод расчётных уровней на график не возможен!")

    def draw_curve(self, morfostvor: Morfostvor, ax: plt.subplot, x="Q", y="УВ"):
        """Отрисовка кривой на графике по заданным из морфоствора параметрам.

        Args:
            morfostvor (Morfostvor): Объект морфоствора из которого получаем данные
            ax (plt.subplot): Ось на которой отрисовывать график
            x (str, optional): Значения по оси x. Defaults to 'Q'.
            y (str, optional): Значения по оси y. Defaults to 'УВ'.
        """

        df = morfostvor.hydraulic_table
        df = df.reset_index(level=0)  # Переводим индекс уровня воды в столбец

        sectors = set(df.index)  # Удаляем дублирующиеся записи
        sectors.remove("Сумма")  # Удаляем запись суммирующего участка

        # Отрисовка суммирующей кривой на графике
        ax.plot(
            df.loc[("Сумма"), x],
            df.loc[("Сумма"), y],
            label="Сумма",
            linewidth=3,
            color="red",
        )

        # Отрисовка кривых по участкам
        for sector in sectors:
            ax.plot(
                df.loc[(sector), x],
                df.loc[(sector), y],
                "--",
                label=sector,
                color=self.sector_colors[sector],
            )

        # Отрисовка легенды
        ax.legend(loc="lower right", fontsize=config.FONT_SIZE["legend"])


@dataclass
class GraphQH(GraphCurve):
    # Номер рисунка
    _fig_num = 2

    # Подписи осей
    _x_label_text = "Q, м³/с"
    _y_label_text = "H, м"
    _ax_title_text = "Гидравлическая кривая"

    def draw(self):
        y_min = min(
            self.morfostvor.hydraulic_table.reset_index(0)
            .loc["Сумма"]
            .reset_index(drop=True)["УВ"]
        )
        self.draw_curve(self.morfostvor, self.ax, "Q", "УВ")
        self.draw_water_levels(self.morfostvor, self.ax, "Q", "H", y_min)


@dataclass
class GraphQV(GraphCurve):
    # Номер рисунка
    _fig_num = 3
    _fig_size = (16.5, 11)
    fig: plt.figure = plt.figure(_fig_num, figsize=_fig_size)
    ax: plt.subplot = fig.add_subplot(111)

    # Подписи осей
    _x_label_text = "Q, м³/с"
    _y_label_text = "V, м/c"
    _ax_title_text = "Кривая скоростей"

    def draw(self):
        self.draw_curve(self.morfostvor, self.ax, "Q", "V")
        self.draw_water_levels(self.morfostvor, self.ax, "Q", "V")


@dataclass
class GraphQF(GraphCurve):
    # Номер рисунка
    _fig_num = 4
    _fig_size = (16.5, 11)
    fig: plt.figure = plt.figure(_fig_num, figsize=_fig_size)
    ax: plt.subplot = fig.add_subplot(111)

    # Подписи осей
    _x_label_text = "Q, м³/с"
    _y_label_text = "F, м²"
    _ax_title_text = "Кривая площадей"

    def draw(self):
        self.draw_curve(self.morfostvor, self.ax, "Q", "F")
        self.draw_water_levels(self.morfostvor, self.ax, "Q", "F")


@dataclass
class GraphProfile(Graph):
    _fig_size = (16.5, 12)
    _fig_num = 1

    fig: plt.figure = plt.figure(_fig_num, figsize=_fig_size)

    __gs = gridspec.GridSpec(80, 3)

    ax_top: plt.subplot = fig.add_subplot(__gs[0, :], frame_on=False)
    ax: plt.subplot = fig.add_subplot(__gs[1:62, :])
    ax_bottom: plt.subplot = fig.add_subplot(__gs[62:, :])
    ax_bottom_overlay: plt.subplot = fig.add_subplot(__gs[62:, :], frame_on=False)
    
    def __post_init__(self):
        self.clean()

        # Настройка параметров графиков и их инициализация
        self.fig.subplots_adjust(bottom=0.08, left=0.08, right=0.9)

        # Добавляем в список границ максимальную и минимальную отметки
        self._y_limits.append(max(self.morfostvor.y))
        self._y_limits.append(min(self.morfostvor.y))

        self._update_limit()
        self.set_style()

        self.draw_profile_footer()
        self.draw_sectors()
        self.draw_profile_bottom()

    def draw_profile_bottom(self):
        """
        Отрисовка дна профиля.

        :return: Отрисовыает дно профиля на графике ax_profile.
        """

        self.ax.plot(
            self.morfostvor.x,
            self.morfostvor.y,
            color=config.COLOR["profile_bottom"],
            linewidth=config.LINE_WIDTH["profile_bottom"],
            linestyle="solid",
        )

    def draw_profile_footer(self):
        """
        Отрисовка подвала с информацией о профиле.

            :param self:
        """
        n = 5  # Количество ячеек подвала
        hs = 10  # Стандартная высота ячейки подвала
        x1 = self.morfostvor.x[0]
        x2 = self.morfostvor.x[-1]

        def setup_box():
            y_top = n * hs

            # Технический разделитель (для увеличения размера границ)
            self.ax_bottom_overlay.plot(
                (x1, x2), (y_top, y_top), alpha=0, color="red"
            )

            self.ax_bottom.plot(
                (x1, x1), (0, y_top), alpha=0, color="red"
            )

        def draw_pk():
            """  Отрисовывает нижнюю границу для ПК в подвале, сами значения ПК отрисовываются отдельно
            """
            y_top = n*hs
            y_bot = n*hs-hs
            y_mid = y_top - ((y_top - y_bot) / 2)
            x2 = self.morfostvor.x[-1]

            # Подпись ячейки
            label = 'ПК'
            self.ax_bottom_overlay.text(
                x2, y_mid, "   " + label,
                color=config.COLOR["bottom_text"],
                fontsize=config.FONT_SIZE["bottom_description"],
                horizontalalignment='left', verticalalignment='center',
            )

            self.ax_bottom_overlay.plot(
                (x1, x2),
                (y_bot, y_bot),
                color=config.COLOR["border"],
                linewidth=config.LINE_WIDTH["profile_bottom"],
                linestyle="solid",
            )

        def draw_h(num=4):
            y_top = num*hs
            y_bot = num*hs-hs
            y_mid = y_top - ((y_top - y_bot) / 2)
            x2 = self.morfostvor.x[-1]

            # Подпись ячейки
            label = 'Отметка'
            self.ax_bottom_overlay.text(
                x2, y_mid, "   " + label,
                color=config.COLOR["bottom_text"],
                fontsize=config.FONT_SIZE["bottom_description"],
                horizontalalignment='left', verticalalignment='center',
            )

            # Верхняя граница
            self.ax_bottom_overlay.plot(
                (x1, x2),
                (y_top, y_top),
                color=config.COLOR["border"],
                linewidth=config.LINE_WIDTH["profile_bottom"],
                linestyle="solid",
            )

            # Нижняя граница
            self.ax_bottom_overlay.plot(
                (x1, x2),
                (y_bot, y_bot),
                color=config.COLOR["border"],
                linewidth=config.LINE_WIDTH["profile_bottom"],
                linestyle="solid",
            )


            # Цикл по всем точкам
            for i in range(len(self.morfostvor.x)):
                x = self.morfostvor.x[i]
                y = self.morfostvor.y[i]

                # Подписи отметок
                self.ax_bottom.text(
                    x,
                    y_mid,
                    f"{y:.2f}",
                    color=config.COLOR["bottom_text"],
                    fontsize=config.FONT_SIZE["bottom_small"],
                    verticalalignment="center",
                    horizontalalignment="center",
                    rotation="90",
                )

        def draw_dist(num=3):
            y_top = num*hs
            y_bot = num*hs-hs
            y_mid = y_top - ((y_top - y_bot) / 2)     

            x1 = self.morfostvor.x[0]
            x2 = self.morfostvor.x[-1]

            # Подпись ячейки
            label = 'Расстояние'
            self.ax_bottom_overlay.text(
                x2, y_mid, "   " + label,
                color=config.COLOR["bottom_text"],
                fontsize=config.FONT_SIZE["bottom_description"],
                horizontalalignment='left', verticalalignment='center',
            )
            
            # Верхняя граница
            self.ax_bottom_overlay.plot(
                (x1, x2),
                (y_top, y_top),
                color=config.COLOR["border"],
                linewidth=config.LINE_WIDTH["profile_bottom"],
                linestyle="solid",
            )

            # Нижняя граница
            self.ax_bottom_overlay.plot(
                (x1, x2),
                (y_bot, y_bot),
                color=config.COLOR["border"],
                linewidth=config.LINE_WIDTH["profile_bottom"],
                linestyle="solid",
            )

            # Цикл по всем точкам
            for i in range(len(self.morfostvor.x)):
                x = self.morfostvor.x[i]

                # Разделители расстояний между точками
                self.ax_bottom.plot(
                    (x, x),
                    (y_bot, y_top),
                    color=config.COLOR["border"],
                    linewidth=config.LINE_WIDTH["profile_bottom"],
                    linestyle="solid",
                )

                # Подписи расстояний между точками
                if i < len(self.morfostvor.x) - 1:
                    x1_ = self.morfostvor.x[i + 1]
                    # Подписи расстояний между точками
                    self.ax_bottom.text(
                        (x + x1_) / 2,
                        y_mid,
                        f"{round(x1_ - x):d}",
                        color=config.COLOR["bottom_text"],
                        fontsize=config.FONT_SIZE["bottom_main"],
                        verticalalignment="center",
                        horizontalalignment="center",
                    )

        def draw_rough(num=2):
            y_top = num*hs
            y_bot = num*hs-hs
            y_mid = y_top - ((y_top - y_bot) / 2)

            x1 = self.morfostvor.x[0]
            x2 = self.morfostvor.x[-1]

            label = 'Коэфф. n'
            self.ax_bottom_overlay.text(
                x2, y_mid, "   " + label,
                color=config.COLOR["bottom_text"],
                fontsize=config.FONT_SIZE["bottom_description"],
                horizontalalignment='left', verticalalignment='center',
            )

            # Верхняя граница
            self.ax_bottom_overlay.plot(
                (x1, x2),
                (y_top, y_top),
                color=config.COLOR["border"],
                linewidth=config.LINE_WIDTH["profile_bottom"],
                linestyle="solid",
            )

            # Нижняя граница
            self.ax_bottom_overlay.plot(
                (x1, x2),
                (y_bot, y_bot),
                color=config.COLOR["border"],
                linewidth=config.LINE_WIDTH["profile_bottom"],
                linestyle="solid",
            )

            # Цикл по участкам
            for sector in self.morfostvor.sectors:
                x = self.morfostvor.x[sector.start_point]
                x1 = self.morfostvor.x[sector.end_point]

                x_mid = x1 - ((x1 - x) / 2)

                # Подписи коэффициентов шероховатости по участкам
                try:
                    self.ax_bottom.text(
                        x_mid,
                        y_mid,
                        f"{sector.roughness:.3f}",
                        color=config.COLOR["bottom_text"],
                        fontsize=config.FONT_SIZE["bottom_main"],
                        verticalalignment="center",
                        horizontalalignment="center",
                    )

                except ValueError:
                    print(
                        "\nОшибка в указании параметров участков (коэффициент шероховатости \
                        или разделение на участки). Проверить данные."
                    )
                    sys.exit(1)

                # Разделители коэффициентов шероховатости
                # Левая граница
                self.ax_bottom.plot(
                    (x, x),
                    (y_bot, y_top),
                    color=config.COLOR["border"],
                    linewidth=config.LINE_WIDTH["profile_bottom"],
                    linestyle="solid",
                )

                # Правая граница
                self.ax_bottom.plot(
                    (x1, x1),
                    (y_bot, y_top),
                    color=config.COLOR["border"],
                    linewidth=config.LINE_WIDTH["profile_bottom"],
                    linestyle="solid",
                )

        def draw_situation(num=1):
            y_top = num*hs
            y_bot = num*hs-hs
            y_mid = y_top - ((y_top - y_bot) / 2)

            x1 = self.morfostvor.x[0]
            x2 = self.morfostvor.x[-1]

            # Подпись ряда
            label = 'Ситуация'
            self.ax_bottom_overlay.text(
                x2, y_mid, "   " + label,
                color=config.COLOR["bottom_text"],
                fontsize=config.FONT_SIZE["bottom_description"],
                horizontalalignment='left', verticalalignment='center',
            )

            # Верхняя граница
            self.ax_bottom_overlay.plot(
                (x1, x2),
                (y_top, y_top),
                color=config.COLOR["border"],
                linewidth=config.LINE_WIDTH["profile_bottom"],
                linestyle="solid",
            )

            # Нижняя граница
            self.ax_bottom_overlay.plot(
                (x1, x2),
                (y_bot, y_bot),
                color=config.COLOR["border"],
                linewidth=config.LINE_WIDTH["profile_bottom"],
                linestyle="solid",
            )

            for sector in self.morfostvor.situation:
                x1 = self.morfostvor.x[sector.start_point]
                x2 = self.morfostvor.x[sector.end_point]
                x_mid = x2 - ((x2 - x1) / 2)

                if sector.type == "УВ":
                    linestyle = 'solid'
                    linewidth = 2

                    self.ax_bottom.add_patch(
                        Rectangle((x1, y_bot), (x2-x1), hs,
                        facecolor = 'deepskyblue',
                        fill=True)
                    )
                else:
                    linestyle = '--'
                    linewidth = 1

                # Подпись в ситуации
                self.ax_bottom.text(
                    x_mid,
                    y_mid,
                    f"{sector.type}",
                    style='italic',
                    color=config.COLOR["bottom_text"],
                    fontsize=config.FONT_SIZE["bottom_medium"],
                    verticalalignment="center",
                    horizontalalignment="center"
                )

                # Левая граница
                self.ax_bottom.plot(
                    (x1, x1),
                    (y_bot, y_top),
                    color=config.COLOR["border"],
                    linewidth=linewidth,
                    linestyle=linestyle,
                )

                # Правая граница
                self.ax_bottom.plot(
                    (x2, x2),
                    (y_bot, y_top),
                    color=config.COLOR["border"],
                    linewidth=linewidth,
                    linestyle=linestyle,
                )

        print(self.morfostvor.situation)
        setup_box()
        draw_pk()
        draw_h(4)
        draw_dist(3)
        draw_rough(2)
        draw_situation(1)
        # draw_dist(1)

    def draw_sectors(self):
        """
        Отрисовка различной информации связанной с участками профиля.

        :param fill: [bool] - заливка полигонов участков на профиле соответствующими цветами
        :param bottom: [bool] - заливка линии дна соответствующими участкам цветами
        :param label: [bool] - отрисовка названий участков, их длин и стрелок обозначающих границы участков
        :return: Отрисовка графической информации по участкам профиля на графике ax_profile.
        """

        h_max = np.floor(max(self.morfostvor.y)) + 1

        for sector in self.morfostvor.sectors:
            points = []

            for i in range(len(sector.coord[0])):
                points.append((sector.coord[0][i], sector.coord[1][i]))

            points.insert(0, (sector.coord[0][0], h_max))
            points.append((sector.coord[0][-1], h_max))

            polygon = matplotlib.patches.Polygon(
                points, alpha=0.04, linestyle="--", label=sector.name
            )
            polygon.set_color(sector.color)

            # Подписи названий и длин участков со стрелками
            if config.PROFILE_SECTOR_LABEL:
                p0 = 1
                p1 = 2
                p3 = 3

                # Расчёт середины участка (для центровки текста)
                cent_x = sector.coord[0][-1] - (
                    (sector.coord[0][-1] - sector.coord[0][0]) / 2
                )

                # Вывод ширины участка
                self.ax_top.text(
                    cent_x,
                    p1,
                    f"{round(sector.length):d} м",
                    color=config.COLOR["sector_text"],
                    verticalalignment="center",
                    horizontalalignment="center",
                    bbox={
                        "facecolor": "white",
                        "edgecolor": "white",
                        "alpha": 1,
                        "pad": 2.5,
                    },
                )

                self.ax_top.text(
                    cent_x,
                    6,
                    sector.name,
                    color=config.COLOR["sector_text"],
                    verticalalignment="center",
                    horizontalalignment="center",
                )

                # Вывод разделителя участков профиля
                self.ax_top.plot(
                    [sector.coord[0][0], sector.coord[0][0]],
                    [p0, p3],
                    color=config.COLOR["sector_line"],
                    linestyle="-",
                    linewidth=config.LINE_WIDTH["sector_line"],
                )  # Горизонтальная слева

                self.ax_top.plot(
                    [sector.coord[0][-1], sector.coord[0][-1]],
                    [p0, p3],
                    color=config.COLOR["sector_line"],
                    linestyle="-",
                    linewidth=config.LINE_WIDTH["sector_line"],
                )  # Горизонтальная справа

                self.ax_top.plot(
                    [sector.coord[0][0], cent_x],
                    [p1, p1],
                    color=config.COLOR["sector_line"],
                    linestyle="-",
                    linewidth=config.LINE_WIDTH["sector_line"],
                )  # Вертикальная слева

                self.ax_top.plot(
                    [cent_x, sector.coord[0][-1]],
                    [p1, p1],
                    color=config.COLOR["sector_line"],
                    linestyle="-",
                    linewidth=config.LINE_WIDTH["sector_line"],
                )  # Вертикальная справа

            # Заливка на профиле участков
            if config.PROFILE_SECTOR_FILL:
                self.ax.add_patch(polygon)

            # Цвет линии дна по участкам
            if config.PROFILE_SECTOR_BOTTOM_LINE:
                self.ax.plot(sector.coord[0], sector.coord[1], "-", color=sector.color)

    def set_style(self):
        # Устанавливаем заголовки графиков
        if config.GRAPHICS_TITLES:
            self.ax.set_title(
                self.morfostvor.title,
                color=config.COLOR["title_text"],
                fontsize=config.FONT_SIZE["title"],
                y=1.1,
            )

        self.ax.set_ylim(self._y_lim)

        # Настраиваем границы и толщину линий границ
        self.ax.spines["top"].set_visible(False)
        self.ax.spines["right"].set_visible(False)
        self.ax.spines["left"].set_linewidth(config.LINE_WIDTH["ax_border"])
        self.ax.spines["bottom"].set_linewidth(config.LINE_WIDTH["ax_border"])

        self.ax_bottom.spines["top"].set_visible(False)
        self.ax_bottom.spines["right"].set_linewidth(config.LINE_WIDTH["ax_border"])
        self.ax_bottom.spines["left"].set_linewidth(config.LINE_WIDTH["ax_border"])
        self.ax_bottom.spines["bottom"].set_linewidth(config.LINE_WIDTH["ax_border"])

        # Устанавливаем отступы в графиках
        self.ax.margins(0.025)
        self.ax_top.margins(0.025)
        self.ax_bottom.margins(0.025, 0)
        self.ax_bottom_overlay.margins(0)

        # Устанавливаем прозрачность заливки фона
        self.ax_top.patch.set_alpha(0)
        self.ax_bottom.patch.set_alpha(0)
        self.ax_bottom_overlay.patch.set_alpha(0)

        # Включаем отображение сетки
        self.ax.grid(True, which="both")

        # Включаем отображение второстепенных засечек на осях
        self.ax.minorticks_on()

        # Устанавливаем параметры засечек на основных осях
        self.ax.tick_params(
            which="major",
            direction="out",
            width=2,
            length=5,
            pad=13,
            labelcolor=config.COLOR["ax_label_text"],
            labelsize=config.FONT_SIZE["ax_major"],
        )

        self.ax.tick_params(
            which="minor",
            direction="out",
            width=1.5,
            length=3,
            pad=10,
            labelcolor=config.COLOR["ax_label_text"],
            labelsize=config.FONT_SIZE["ax_minor"],
        )

        # Отключаем засечки и подписи на осях вспомогательных графиков
        self.ax_bottom.set_xticks([])
        self.ax_bottom.set_yticks([])
        self.ax_bottom_overlay.set_xticks([])
        self.ax_bottom_overlay.set_yticks([])
        self.ax_top.set_xticks([])
        self.ax_top.set_yticks([])

        # Устанавливаем параметры подписей осей
        self.ax.set_ylabel(
            "H, м",
            color=config.COLOR["ax_label_text"],
            fontsize=config.FONT_SIZE["ax_label"],
            fontstyle="italic",
        )

        self.ax.yaxis.set_label_coords(-0.025, 1.08)

        # Устанавливает параметры вывода значений осей
        self.ax.yaxis.set_major_formatter(matplotlib.ticker.FormatStrFormatter("%.10g"))

        # Настройка параметров отображение сетки
        self.ax.grid(
            which="major",
            color=config.COLOR["ax_grid"],
            linestyle=":",
            linewidth=1,
            alpha=0.9,
        )

        self.ax.grid(
            which="minor",
            color=config.COLOR["ax_grid_sub"],
            linestyle=":",
            linewidth=1,
            alpha=0.9,
        )

    def draw_profile_point_lines(self):
        """
        Отрисовка вертикальных линий от точек до подвала.

        """
        for i in range(len(self.morfostvor.x)):
            self.ax.plot(
                (self.morfostvor.x[i], self.morfostvor.x[i]),
                (self.morfostvor.y[i], self._y_lim[0]),
                color=config.COLOR["profile_point_line"],
                linewidth=config.LINE_WIDTH["profile_point_line"],
                linestyle="solid",
            )

    def draw_erosion_limit(self, h, x1=None, x2=None, text="▼$H_{{разм.}} = {h:.2f}$"):
        """Функция отрисовки линии предельного размыва профиля.

        Arguments:
            h {[float]} -- Отметка линии предельного размыва

        Keyword Arguments:
            x1 {[float]} -- Координата начала линии (default: {None})
            x2 {[float]} -- Координата конца линии (default: {None})
            text {[string]} -- Текст подписи линии (default: {'▼$H_{{разм.}} = {h:.2f}$'})
        """
        if config.PROFILE_EROSION_LIMIT and not isinstance(
            self.morfostvor.erosion_limit, str
        ):
            # Ограничение линии предельного размыва
            # по всему профилю если параметр config.PROFILE_EROSION_LIMIT_FULL = true
            if config.PROFILE_EROSION_LIMIT_FULL:
                x1 = min(self.morfostvor.x)
                x2 = max(self.morfostvor.x)
            # Если координаты начала и конца линии не заданы, устанавливаем по границе профиля
            # если есть участки 'Левая пойма', 'Правая пойма' задаем границы линии по участкам
            else:
                if x1 is None:
                    x1 = min(self.morfostvor.x)
                    for sector in self.morfostvor.sectors:
                        if sector.name == "Левая пойма":
                            x1 = sector.coord[0][-1]
                if x2 is None:
                    x2 = max(self.morfostvor.x)
                    for sector in self.morfostvor.sectors:
                        if sector.name == "Правая пойма":
                            x2 = sector.coord[0][0]

            # Подпись текста
            erosion_limit_text = self.ax.text(
                x2 - 1,
                h + 0.01,
                text.format(h=h),
                color=config.COLOR["erosion_limit_text"],
                fontsize=config.FONT_SIZE["erosion_limit"],
                weight="bold",
            )
            # Обводка текста
            erosion_limit_text.set_path_effects(
                [
                    path_effects.Stroke(linewidth=3, foreground="white", alpha=0.95),
                    path_effects.Normal(),
                ]
            )

            # Отрисовка линии предельного размыва
            self.ax.plot(
                [x1, x2],
                [h, h],
                color=config.COLOR["erosion_limit_line"],
                linestyle="--",
                linewidth=config.LINE_WIDTH["erosion_limit_line"],
            )
            # Добавляем в список границ отметку
            self._y_limits.append(h)
            self._update_limit()

    def draw_top_limit(self, h, x1=None, x2=None, text="{}\nH = {:.2f}"):
        # Если координаты начала и конца линии не заданы, устанавливаем по границе профиля
        # если есть участки 'Левая пойма', 'Правая пойма' задаем границы линии по участкам
        if x1 is None:
            x1 = min(self.morfostvor.x)
            for sector in self.morfostvor.sectors:
                if sector.name == "Левая пойма":
                    x1 = sector.coord[0][-1]
        if x2 is None:
            x2 = max(self.morfostvor.x)
            for sector in self.morfostvor.sectors:
                if sector.name == "Правая пойма":
                    x2 = sector.coord[0][0]
        y_step = self.ax.get_yticks()[1] - self.ax.get_yticks()[0]
        cent_x = x2 - ((x2 - x1) / 2)

        top_limit_text = self.ax.text(
            cent_x,
            h + (y_step * 0.2),
            f"{self.morfostvor.top_limit_description}\nH = {h:.2f}",
            color=config.COLOR['top_limit_text'],
            fontsize=config.FONT_SIZE['top_limit'],
            weight='bold',
            horizontalalignment='center',
            verticalalignment='center')

        self.ax.plot(
            [x1, x2],
            [h, h],
            color=config.COLOR["top_limit_line"],
            linestyle="-.",
            linewidth=config.LINE_WIDTH["top_limit_line"],
        )

        self._y_limits.append(h)
        self._update_limit()

    def draw_waterline(
        self,
        h,
        color=config.COLOR["water_line"],
        linestyle="--",
        linewidth=config.LINE_WIDTH["water_line"],
    ):
        """
        Функция отрисовки уреза воды по границам водного объекта.

        :param water: Исходный водный объект, содержащий координаты границ воды.
        :return: урез на графике профиля (ax_profile).
        """

        def draw_line(self):
            for boundary in water.boundary():
                # Вводим служебные координаты
                x1, x2 = boundary[0][0], boundary[0][1]  # Начало и конец x
                y1, y2 = boundary[1][0], boundary[1][1]  # отметки уреза

                # Рисуем урез воды
                self.ax.plot(
                    [x1, x2],
                    [y1, y2],
                    color=color,
                    linestyle=linestyle,
                    linewidth=linewidth,
                )

                if config.PROFILE_WATER_FILL:
                    self.ax.fill(
                        water.water_section_x,
                        water.water_section_y,
                        facecolor=config.COLOR["water_fill"],
                        alpha=0.2,
                    )

        if config.OVERFLOW:
            water = WaterSection(self.morfostvor.x, self.morfostvor.y, h)
            draw_line(self)

        else:
            # Рисуем урезы на каждом участке
            for sector in self.morfostvor.sectors:
                x = sector.coord[0]
                y = sector.coord[1]

                if h >= min(y):
                    water = WaterSection(x, y, h)
                    draw_line(self)

        self._update_limit()
        self.set_style()

    def draw_levels_on_profile(self, levels):
        """
        Функция отрисовки полученных расчётных уровней воды на поперечном профиле.

        :param levels: DataFrame содержащий столбцы P, Q, H
        :return:
        """
        label = []

        for index, row in levels.iterrows():
            # Отрисовка уреза
            water_level = row["H"]

            self.draw_waterline(water_level)

            if config.PROFILE_LEVELS_TITLE:
                # Подпись уровня воды на профиле
                water = WaterSection(self.morfostvor.x, self.morfostvor.y, water_level)
                try:
                    water = WaterSection(
                        self.morfostvor.x, self.morfostvor.y, water_level
                    )
                except:
                    print("Ошибка! При отрисовке расчётных уровней на профиле. \n")

                padding = 0.01
                x = water.water_section_x[0] + 2 * padding
                y = water_level + padding

                try:
                    # Если обеспеченность записана цифрами
                    waterline_text = self.ax.text(
                        x,
                        y,
                        f"▼$P_{{{row['P']:2g}\\%}} = {row['H']:.2f}$",
                        color=config.COLOR["water_level_text"],
                        fontsize=config.FONT_SIZE["water_level"],
                        weight="bold",
                    )
                    waterline_text.set_path_effects(
                        [
                            path_effects.Stroke(
                                linewidth=3, foreground="white", alpha=0.55
                            ),
                            path_effects.Normal(),
                        ]
                    )
                except ValueError:
                    # Если обеспеченность записана строкой
                    waterline_text = self.ax.text(
                        x,
                        y,
                        f"{row['P']} = {row['H']:.2f}",
                        color=config.COLOR["water_level_text"],
                        fontsize=config.FONT_SIZE["water_level"],
                        weight="bold",
                    )

                    waterline_text.set_path_effects(
                        [
                            path_effects.Stroke(
                                linewidth=1.8, foreground="white", alpha=0.55
                            ),
                            path_effects.Normal(),
                        ]
                    )

            try:
                label.append(f"$P_{{{row['P']:2g}\\%}} = {water_level:.2f}$ м\n")
            except ValueError:
                label.append(f"${row['P']} = {water_level:.2f}$ м\n")

            # Вывод линий сносок от уровней воды к таблице
            if config.PROFILE_LEVELS_TABLE_LINES:
                water = WaterSection(self.morfostvor.x, self.morfostvor.y, water_level)

                # Горизонтальные точки линий сносок
                x_step = (water.water_section_x[-1] - water.water_section_x[0]) / len(
                    self.morfostvor.probability
                )
                # Нижняя координата x
                x0 = water.water_section_x[0] + (x_step * (index + 1) / 2)
                x1 = x0 + (x0 / 8 * (index + 1))  # Верхняя координата x
                x_lim = self.ax.get_xlim()  # Получаем границы графика
                x3 = x_lim[1]  # Координата x границы справа
                self.ax.set_xlim(x_lim)  # Возвращаем границы на исходные

                # Вертикальные точки линий сносок
                # 1% вертикальный от графика
                y_step = (self.top_limit - self.bottom_limit) / 100
                y0 = water_level  # Нижняя координата y (отметка уреза воды)
                if index == 0:
                    # Верхняя координата y для первой линии уреза
                    y1 = self.top_limit - (y_step) - (y_step * 3 * (index))
                else:
                    # Верхняя координата y для последующих линий уреза
                    y1 = self.top_limit - (y_step * 2.95 * (index))

                # Устанавливаем параметры отображения линий сносок
                color = config.COLOR["water_reference_line"]
                linestyle = "--"
                linewidth = config.LINE_WIDTH["water_line"] / 1.75
                alpha = 0.8

                # Линии сносок
                self.ax.plot(
                    [x0, x1],
                    [y0, y1],
                    color=color,
                    linestyle=linestyle,
                    linewidth=linewidth,
                    alpha=alpha,
                )
                self.ax.plot(
                    [x1, x3],
                    [y1, y1],
                    color=color,
                    linestyle=linestyle,
                    linewidth=linewidth,
                    alpha=alpha,
                )

        if self.morfostvor.waterline and type(self.morfostvor.waterline) is not str:
            label.append(f"\nУВ = {self.morfostvor.waterline:.2f} м\n")

            if self.morfostvor.date:
                label.append(f"({self.morfostvor.date})")

        if config.PROFILE_WATER_LEVEL_NOTE:
            if self.morfostvor.waterline == "-" or self.morfostvor.waterline == "":
                label.append("\nПримечание: на\nмомент съёмки\nсток отсутствует")

        # Вывод таблицы уровней с разными обеспеченностями (справа)
        self.ax.annotate(
            "".join(label).rstrip(),
            xy=(1, 1),
            ha="left",
            va="top",
            xycoords="axes fraction",
            size=config.FONT_SIZE["levels_table"],
            color=config.COLOR["levels_table"],
            bbox=dict(boxstyle="round", fc="white", ec="none"),
        )

    def draw_wet_perimeter(self):
        """Функция отрисовки смоченного периметра на графике поперечного профиля"""

        # Проверяем задан ли расчётный шаг в исходных данных
        if isinstance(self.morfostvor.dH, str) or self.morfostvor.dH == 0:
            self.morfostvor.dH = 1
            dH = self.morfostvor.dH
        else:
            dH = self.morfostvor.dH

        # Переводим сантиметры приращения в метры
        dH = dH / 100

        min_sector = self.morfostvor.get_min_sector()

        # Исходные сектора для расчёта (сектор, содержащий минимальную отметку)
        calc_sectors = [min_sector[0]]

        # Уровень воды, с минимальным отступом
        water_level = min(self.morfostvor.y) + dH

        # Цикл расчёта до максимального уровня воды
        while water_level < self.morfostvor.hydraulic_result["УВ"].max():
            if config.OVERFLOW:
                for i in calc_sectors:
                    sector = self.morfostvor.sectors[i]
                    x = sector.coord[0]
                    y = sector.coord[1]

                    # Максимальная отметка слева
                    previous_min_ele = max(chunk_list(y, 2)[0])
                    # Максимальная отметка справа
                    next_min_ele = max(chunk_list(y, 2)[1])

                    # Проверка на перелив через границы участка
                    if (
                        (water_level >= previous_min_ele)
                        and (i - 1 not in calc_sectors)
                        and (i - 1 >= 0)
                    ):
                        calc_sectors.append(i - 1)
                    if (
                        (water_level >= next_min_ele)
                        and (i + 1 not in calc_sectors)
                        and (i + 1 <= len(self.morfostvor.sectors) - 1)
                    ):
                        calc_sectors.append(i + 1)

                    # Сектор воды и основные его параметры
                    # Расчетный участок является участком с минимальными отметками
                    # либо расчёт выполняется с одновременным заполнением
                    # начинаем заполнять с точки с минимальной отметкой
                    if sector.id == min_sector[1].id:
                        water = WaterSection(x, y, water_level)

                    # Расчетный участок находится слева от начального
                    # начинаем заполнять с крайней правой точки
                    elif sector.id < min_sector[1].id:
                        water = WaterSection(
                            x, y, water_level, start_point=[len(y) - 1, y[-1]]
                        )

                    # Расчетный участок находится справа от начального
                    # начинаем заполнять с крайней левой точки
                    elif sector.id > min_sector[1].id:
                        water = WaterSection(x, y, water_level, start_point=[0, y[0]])

                    # Отрисовка смоченного периметра на профиле на профиле
                    self.ax.plot(
                        water.water_section_x,
                        water.water_section_y,
                        ":",
                        marker="o",
                        linewidth=1,
                        color="black",
                        markersize=3,
                    )
                    self.ax.plot(
                        [water.water_section_x[0], water.water_section_x[-1]],
                        [water.water_section_y[0], water.water_section_y[-1]],
                        ":",
                        linewidth=1,
                        color="black",
                    )
            else:
                # Отрисовка с заполнением по участкам
                for sector in self.morfostvor.sectors:
                    x = sector.coord[0]
                    y = sector.coord[1]

                    if min(y) < water_level:
                        # Сектор воды и основные его параметры
                        water = WaterSection(x, y, water_level)

                        # Отрисовка смоченного периметра на профиле
                        self.ax.plot(
                            water.water_section_x,
                            water.water_section_y,
                            ":",
                            marker="o",
                            linewidth=1,
                            color="black",
                            markersize=3,
                        )
                        self.ax.plot(
                            [water.water_section_x[0], water.water_section_x[-1]],
                            [water.water_section_y[0], water.water_section_y[-1]],
                            ":",
                            linewidth=1,
                            color="black",
                        )

            water_level += dH

    def _update_limit(self):
        # Шаг засечек по вертикали
        y_step = self.ax.get_yticks()[1] - self.ax.get_yticks()[0]

        # Минимальное и максимальное значения из списка границ
        min_y = min(self._y_limits)
        max_y = max(self._y_limits)

        # Нижняя граница
        self.bottom_limit = np.ceil(min_y) - y_step
        if self.morfostvor.erosion_limit:
            self.bottom_limit = np.ceil(self.morfostvor.erosion_limit) - y_step
            while (self.morfostvor.erosion_limit - self.bottom_limit) < (y_step / 3):
                self.bottom_limit -= y_step
        else:
            self.bottom_limit = np.ceil(min_y) - y_step

            while (self.morfostvor.ele_min - self.bottom_limit) < (y_step / 3):
                self.bottom_limit -= y_step

        # Верхняя граница
        if y_step > 0.5:
            self.top_limit = round(np.floor(max_y) + y_step, 3)
        else:
            self.top_limit = round((max_y // y_step * y_step) + y_step * 2, 3)

        # Устанавливаем границы отображения
        self._y_lim = (self.bottom_limit, self.top_limit)
        self.ax.set_ylim(self._y_lim)
        self.draw_profile_point_lines()


def xls_calculate_hydraulic(in_filename, out_filename, page=None):
    """
    Выполнение гидравлических расчетов и создание отчета по результатам расчетов.
    Исходные данные представлены в in_filename (xls файл). По умолчанию расчеты производятся
    для всех листов xls файла. Если задан параметр page, расчет производится только для указанной страницы.
    По результат создается out_filename (результирующий отчет в формате docx).

        :param in_filename: Входные данные по створам (.xls или .xlsx файл)
        :param out_filename: Результаты расчетов  (.docx файл)
        :param page=None: Номер страницы в xls файле, по умолчанию None (расчеты производятся для всего документа)
    """

    # Создаем родительскую папку, если она не существует
    Path(out_filename).parents[0].mkdir(parents=True, exist_ok=True)

    # Удаляем предыдущий отчет, если включена перезапись файла
    if config.REWRITE_DOC_FILE:
        try:
            os.remove(out_filename)
        except:
            pass

    page_quantity = get_xls_sheet_quantity(in_filename)
    stvors = []

    # Расчет для всех листов xls файла
    if page is None:
        for i in range(page_quantity):
            stvors.append(Morfostvor())
            stvors[i].read_xls(in_filename, i)
            stvors[i].calculate()
            stvors[i].doc_export(out_filename)

        # Вставка сводных таблиц
        insert_summary_QV_tables(stvors, out_filename)

    # Расчет только одного листа xls файла
    elif type(page) == int:
        stvor = Morfostvor()
        stvor.read_xls(in_filename, page)
        stvor.calculate()
        stvor.doc_export(out_filename)
    else:
        print("Номер листа должен быть int.")
        sys.exit(0)
